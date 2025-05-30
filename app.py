import os
from datetime import datetime, timedelta, date, timezone
import json
from flask import Flask, render_template, url_for, flash, redirect, request, session, jsonify, send_file, send_from_directory, g
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, current_user, logout_user, login_required
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField, BooleanField, SelectField, FileField
from wtforms.validators import DataRequired, Length, Email, EqualTo, ValidationError
from flask_wtf.file import FileAllowed
from functools import wraps
import traceback
import random
import time
import requests
import time
import urllib.parse
from bs4 import BeautifulSoup
import qrcode
from PIL import Image
import io
import re
import uuid
import pandas as pd
import numpy as np
from telegram_web_scraper import parse_telegram_channel, save_messages_to_json
import ipaddress
import docx
from flask_session import Session
from create_folders import create_necessary_folders
import secrets
import math
from urllib.parse import urljoin
from flask_bcrypt import Bcrypt
import shutil
from PIL import ImageDraw
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# API ключ для Checko.ru
API_KEY = "ZXHNsb6ZeNjtCsRk"
COMPANY_API_KEY = API_KEY  # Для обратной совместимости
# Включаем режим отладки для вывода дополнительной информации
debug_mode = True

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['ALLOWED_EXTENSIONS'] = {'csv', 'xls', 'xlsx'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB
app.config['PROFILE_PICS_FOLDER'] = os.path.join(app.config['UPLOAD_FOLDER'], 'profile_pics')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'flask_session')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_USE_SIGNER'] = True
# Убираем прокси, т.к. используем ручной ввод CAPTCHA
# app.config['TELEGRAM_PROXY_URL'] = os.environ.get('TELEGRAM_PROXY_URL')

# Создаем необходимые директории
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)
os.makedirs(app.config['PROFILE_PICS_FOLDER'], exist_ok=True)

# Настройка секретного ключа
app.config['SECRET_KEY'] = secrets.token_hex(16)
# Настройка базы данных SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Настройка логгера приложения
app.logger.setLevel(logging.DEBUG if app.debug else logging.INFO)

# Инициализация сессий на основе файловой системы
Session(app)

# Инициализация расширений
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Пожалуйста, войдите в систему для доступа к этой странице.'
login_manager.login_message_category = 'info'

# Создание папки для фотографий профиля
PROFILE_PICS_DIR = os.path.join(app.config['UPLOAD_FOLDER'], 'profile_pics')
os.makedirs(PROFILE_PICS_DIR, exist_ok=True)

# Создание стандартного изображения профиля, если его нет
DEFAULT_PROFILE_PIC = os.path.join(PROFILE_PICS_DIR, 'default.jpg')
# Путь к стандартному аватару из static
AVATAR_IMAGE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'img', 'avatar.png')

# Если стандартного изображения профиля нет, копируем avatar.png
if not os.path.exists(DEFAULT_PROFILE_PIC):
    try:
        # Проверяем, существует ли avatar.png
        if os.path.exists(AVATAR_IMAGE):
            # Создаем директорию для профилей, если её нет
            os.makedirs(os.path.dirname(DEFAULT_PROFILE_PIC), exist_ok=True)
            # Копируем avatar.png как стандартное изображение
            shutil.copy(AVATAR_IMAGE, DEFAULT_PROFILE_PIC)
            print(f"Скопирован аватар из {AVATAR_IMAGE} в {DEFAULT_PROFILE_PIC}")
        else:
            # Если avatar.png не найден, создаем простое изображение
            print(f"Файл {AVATAR_IMAGE} не найден, создаю стандартное изображение")
            img = Image.new('RGB', (200, 200), color=(13, 110, 253))
            d = ImageDraw.Draw(img)
            d.ellipse((50, 50, 150, 150), fill=(255, 255, 255))
            d.ellipse((70, 70, 130, 130), fill=(13, 110, 253))
            img.save(DEFAULT_PROFILE_PIC)
    except Exception as e:
        print(f"Ошибка при создании стандартного изображения: {str(e)}")
        # Создаем изображение 200x200 с синим фоном
        img = Image.new('RGB', (200, 200), color=(13, 110, 253))
        d = ImageDraw.Draw(img)
        d.ellipse((50, 50, 150, 150), fill=(255, 255, 255))
        d.ellipse((70, 70, 130, 130), fill=(13, 110, 253))
        img.save(DEFAULT_PROFILE_PIC)

# Модель пользователя
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    department = db.Column(db.String(100), nullable=False)
    position = db.Column(db.String(100), nullable=False)
    profile_image = db.Column(db.String(100), nullable=False, default='avatar.png')
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    is_admin = db.Column(db.Boolean, nullable=False, default=False)
    last_ip = db.Column(db.String(50), nullable=True)
    last_login = db.Column(db.DateTime, nullable=True)
    favorites = db.relationship('UserFavorite', backref='user', lazy=True)
    tz_helper_chats = db.relationship('TzHelperChat', backref='user', lazy=True)
    
    def __repr__(self):
        return f"User('{self.email}', '{self.first_name}', '{self.last_name}')"

# Модель журнала действий пользователя
class UserLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user_email = db.Column(db.String(120), nullable=False)
    action = db.Column(db.String(200), nullable=False)
    ip_address = db.Column(db.String(50), nullable=True)
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    def __repr__(self):
        return f"UserLog('{self.user_email}', '{self.action}', '{self.timestamp}')"

# Модель кэша новостей
class NewsCache(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    link = db.Column(db.String(500), nullable=False)
    date = db.Column(db.String(50), nullable=False)
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    def __repr__(self):
        return f"NewsCache('{self.title}', '{self.date}')"

    @staticmethod
    def to_json(news_items):
        """Преобразует список объектов NewsCache в JSON-представление"""
        return [{'title': item.title, 'link': item.link, 'date': item.date} for item in news_items]

# Модель доступа к функциям
class FeatureAccess(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    target_type = db.Column(db.String(20), nullable=False)  # 'user' или 'department'
    target_id = db.Column(db.String(100), nullable=False)   # id пользователя или название отдела
    feature = db.Column(db.String(50), nullable=False)
    access = db.Column(db.Boolean, nullable=False, default=True)
    
    def __repr__(self):
        return f"FeatureAccess('{self.target_type}', '{self.target_id}', '{self.feature}', '{self.access}')"

# Новая модель для хранения избранных каналов пользователя
class UserFavorite(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    channel = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    
    def __repr__(self):
        return f"UserFavorite('{self.user_id}', '{self.channel}')"

# Модель для хранения сессий чата ТЗhelper
class TzHelperChat(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False, default='Новое ТЗ')
    content = db.Column(db.Text, nullable=True)  # Сохраненное ТЗ
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow, onupdate=datetime.utcnow)
    chat_history = db.relationship('TzHelperMessage', backref='chat', lazy=True, cascade='all, delete-orphan')
    
    def __repr__(self):
        return f"TzHelperChat('{self.user_id}', '{self.title}', '{self.created_at}')"
        
# Модель для хранения сообщений в чате ТЗhelper
class TzHelperMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    chat_id = db.Column(db.Integer, db.ForeignKey('tz_helper_chat.id'), nullable=False)
    role = db.Column(db.String(20), nullable=False)  # 'user' или 'assistant'
    content = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    
    def __repr__(self):
        return f"TzHelperMessage('{self.chat_id}', '{self.role}', '{self.timestamp}')"

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.before_request
def before_request():
    """
    Проверяет перед каждым запросом, соответствует ли текущий IP адрес последнему известному IP пользователя
    """
    # Если пользователь аутентифицирован
    if current_user.is_authenticated:
        # Обновляем время последней активности
        session.modified = True
        
        # Проверка IP-адреса (только для защищенных маршрутов)
        if request.endpoint and request.endpoint not in ['static', 'logout', 'login', 'register']:
            if hasattr(current_user, 'last_ip') and current_user.last_ip:
                # Если IP значительно изменился и это не первый вход
                if current_user.last_ip != request.remote_addr:
                    # Записываем новый IP
                    current_user.last_ip = request.remote_addr
                    current_user.last_login = datetime.utcnow()
                    
                    # Логируем изменение IP
                    try:
                        log = UserLog(
                            user_id=current_user.id,
                            user_email=current_user.email,
                            action=f"Изменение IP адреса: {current_user.last_ip} → {request.remote_addr}",
                            ip_address=request.remote_addr
                        )
                        db.session.add(log)
                        db.session.commit()
                    except:
                        db.session.rollback()
            else:
                # Если это первый вход, просто фиксируем IP
                current_user.last_ip = request.remote_addr
                current_user.last_login = datetime.utcnow()
                try:
                    db.session.commit()
                except:
                    db.session.rollback()

# Формы для регистрации и входа
class RegistrationForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Пароль', validators=[DataRequired(), Length(min=8)])
    confirm_password = PasswordField('Подтвердите пароль', validators=[DataRequired(), EqualTo('password')])
    first_name = StringField('Имя', validators=[DataRequired()])
    last_name = StringField('Фамилия', validators=[DataRequired()])
    department = SelectField('Отдел', validators=[DataRequired()], 
                           choices=[
                               ('Руководство', 'Руководство'),
                               ('Управление делами', 'Управление делами'),
                               ('Служба безопасности', 'Служба безопасности'),
                               ('Служба ИТ', 'Служба ИТ'),
                               ('Служба сопровождения 1С', 'Служба сопровождения 1С'),
                               ('Юридический отдел', 'Юридический отдел'),
                               ('Транспортная служба', 'Транспортная служба'),
                               ('Финансовая служба', 'Финансовая служба'),
                               ('Отдел маркетинга', 'Отдел маркетинга'),
                               ('Отдел продаж БМИ', 'Отдел продаж БМИ'),
                               ('Конструкторское бюро', 'Конструкторское бюро'),
                               ('Служба директора по сервисному обслуживанию', 'Служба директора по сервисному обслуживанию'),
                               ('Служба качества', 'Служба качества'),
                               ('Служба развития', 'Служба развития'),
                               ('Служба технического директора', 'Служба технического директора'),
                               ('Сервисная служба', 'Сервисная служба'),
                               ('Метрологическая служба', 'Метрологическая служба'),
                               ('Отдел персонала', 'Отдел персонала')
                           ])
    position = StringField('Должность', validators=[DataRequired()])
    submit = SubmitField('Регистрация')

    def validate_email(self, email):
        user = User.query.filter_by(email=email.data).first()
        if user:
            raise ValidationError('Этот email уже зарегистрирован. Пожалуйста, выберите другой.')

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Пароль', validators=[DataRequired()])
    remember = BooleanField('Запомнить меня')
    submit = SubmitField('Войти')

class UpdateProfileForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    first_name = StringField('Имя', validators=[DataRequired()])
    last_name = StringField('Фамилия', validators=[DataRequired()])
    department = SelectField('Отдел', validators=[DataRequired()], 
                           choices=[
                               ('Руководство', 'Руководство'),
                               ('Управление делами', 'Управление делами'),
                               ('Служба безопасности', 'Служба безопасности'),
                               ('Служба ИТ', 'Служба ИТ'),
                               ('Служба сопровождения 1С', 'Служба сопровождения 1С'),
                               ('Юридический отдел', 'Юридический отдел'),
                               ('Транспортная служба', 'Транспортная служба'),
                               ('Финансовая служба', 'Финансовая служба'),
                               ('Отдел маркетинга', 'Отдел маркетинга'),
                               ('Отдел продаж БМИ', 'Отдел продаж БМИ'),
                               ('Конструкторское бюро', 'Конструкторское бюро'),
                               ('Служба директора по сервисному обслуживанию', 'Служба директора по сервисному обслуживанию'),
                               ('Служба качества', 'Служба качества'),
                               ('Служба развития', 'Служба развития'),
                               ('Служба технического директора', 'Служба технического директора'),
                               ('Сервисная служба', 'Сервисная служба'),
                               ('Метрологическая служба', 'Метрологическая служба'),
                               ('Отдел персонала', 'Отдел персонала')
                           ])
    position = StringField('Должность', validators=[DataRequired()])
    picture = FileField('Обновить фото профиля', validators=[FileAllowed(['jpg', 'png', 'jpeg'])])
    submit = SubmitField('Обновить')

    def validate_email(self, email):
        # Проверяем, что текущий пользователь аутентифицирован
        if current_user.is_authenticated:
            # Если email не изменился, валидация не нужна
            if email.data == current_user.email:
                return
            
            # Проверяем, занят ли уже этот email другим пользователем
            user = User.query.filter_by(email=email.data).first()
            if user and user.id != current_user.id:
                raise ValidationError('Этот email уже зарегистрирован. Пожалуйста, выберите другой.')

class UpdatePasswordForm(FlaskForm):
    current_password = PasswordField('Текущий пароль', validators=[DataRequired()])
    new_password = PasswordField('Новый пароль', validators=[DataRequired(), Length(min=8)])
    confirm_password = PasswordField('Подтвердите новый пароль', validators=[DataRequired(), EqualTo('new_password')])
    submit = SubmitField('Обновить пароль')

# Функция для сохранения фотографии профиля
def save_picture(form_picture):
    try:
        print(f"Сохранение изображения профиля: {form_picture.filename}")
        random_hex = secrets.token_hex(8)
        _, f_ext = os.path.splitext(form_picture.filename)
        picture_fn = random_hex + f_ext
        picture_path = os.path.join(PROFILE_PICS_DIR, picture_fn)
        
        # Изменяем размер изображения
        output_size = (200, 200)
        
        try:
            # Открываем изображение
            img = Image.open(form_picture)
            print(f"Изображение открыто: формат={img.format}, размер={img.size}, режим={img.mode}")
            
            # Сохраняем исходное изображение в режиме RGB (если это не RGB)
            if img.mode != 'RGB':
                print(f"Конвертация из {img.mode} в RGB")
                img = img.convert('RGB')
            
            # Создаем миниатюру
            img.thumbnail(output_size)
            print(f"Создана миниатюра: размер={img.size}")
            
            # Создаем квадратное изображение с белым фоном
            background = Image.new('RGB', output_size, (255, 255, 255))
            
            # Вычисляем положение миниатюры на фоне (центрирование)
            offset = ((output_size[0] - img.width) // 2, (output_size[1] - img.height) // 2)
            
            # Накладываем миниатюру на фон
            background.paste(img, offset)
            
            # Сохраняем изображение
            print(f"Сохранение изображения: {picture_path}")
            background.save(picture_path)
            
            return picture_fn
            
        except Exception as e:
            print(f"Ошибка при обработке изображения: {str(e)}")
            # Если произошла ошибка при обработке изображения, сохраняем оригинал
            form_picture.save(picture_path)
            return picture_fn
            
    except Exception as e:
        print(f"Критическая ошибка при сохранении изображения: {str(e)}")
        # Если не удалось сохранить, возвращаем стандартное изображение
        return 'default.jpg'

# Функция для логирования действий пользователя
def log_user_action(action):
    if current_user.is_authenticated:
        log = UserLog(
            user_id=current_user.id,
            user_email=current_user.email,
            action=action,
            ip_address=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()

# Декоратор для проверки прав администратора
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('У вас нет доступа к этой странице.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

# Проверка доступа к функциям
def has_feature_access(user, feature):
    # Если пользователь администратор - полный доступ
    if user.is_admin:
        return True
    
    # Специальная проверка для дашбордов - только для руководителей
    if feature == 'dashboard' and 'руководитель' not in user.position.lower():
        return False
    
    # Проверяем индивидуальные настройки пользователя
    user_access = FeatureAccess.query.filter_by(
        target_type='user',
        target_id=str(user.id),
        feature=feature
    ).first()
    
    if user_access:
        return user_access.access
    
    # Проверяем настройки отдела
    department_access = FeatureAccess.query.filter_by(
        target_type='department',
        target_id=user.department,
        feature=feature
    ).first()
    
    if department_access:
        return department_access.access
    
    # По умолчанию доступ разрешен
    return True

@app.route('/')
def index():
    return render_template('index.html', has_feature_access=has_feature_access)

@app.route('/qr')
@login_required
def qr_generator():
    if not has_feature_access(current_user, 'qr'):
        flash('У вас нет доступа к этой функции.', 'danger')
        return redirect(url_for('index'))
    return render_template('qr.html')

@app.route('/ai')
@login_required
def ai():
    # Проверяем доступ к функции
    if not has_feature_access(current_user, 'ai'):
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('index'))
    return render_template('ai.html')

@app.route('/holdings')
@login_required
def holdings():
    if not has_feature_access(current_user, 'holdings'):
        flash('У вас нет доступа к этой функции.', 'danger')
        return redirect(url_for('index'))
    return render_template('holdings.html')

@app.route('/analytics')
@login_required
def analytics():
    # Проверяем доступ к функции
    if not has_feature_access(current_user, 'analytics'):
        flash('У вас нет доступа к этой функции. Обратитесь к администратору.', 'warning')
        return redirect(url_for('index'))
    return render_template('analytics.html')

@app.route('/competitors')
@login_required
def competitors():
    # Проверяем доступ к функции
    if not has_feature_access(current_user, 'competitors'):
        flash('У вас нет доступа к этой функции. Обратитесь к администратору.', 'warning')
        return redirect(url_for('index'))
    return render_template('competitors.html')

@app.route('/dashboard')
@login_required
def dashboard():
    # Проверяем доступ к функции
    if not current_user.is_admin and 'руководитель' not in current_user.position.lower():
        flash('У вас нет доступа к этой странице.', 'danger')
        return redirect(url_for('index'))
    
    # Получаем список сотрудников отдела
    department_employees = User.query.filter_by(department=current_user.department).order_by(User.last_name).all()
    department_count = len(department_employees)
    
    return render_template('dashboard.html', 
                         department_employees=department_employees,
                         department_count=department_count)

@app.route('/admin')
@login_required
@admin_required
def admin():
    # Получаем всех пользователей
    users = User.query.all()
    
    # Получаем уникальные отделы для управления доступом
    departments = db.session.query(User.department).distinct().all()
    departments = [d[0] for d in departments]
    
    # Получаем последние 50 записей журнала
    logs = UserLog.query.order_by(UserLog.timestamp.desc()).limit(50).all()
    
    # Логируем действие
    log_user_action("Посещение панели администратора")
    
    return render_template('admin.html', users=users, departments=departments, logs=logs)

@app.route('/api/admin/user/<int:user_id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
@admin_required
def admin_user(user_id):
    user = User.query.get_or_404(user_id)
    
    # GET-запрос: получение данных пользователя
    if request.method == 'GET':
        return jsonify({
            'id': user.id,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'department': user.department,
            'position': user.position,
            'is_admin': user.is_admin
        })
    
    # PUT-запрос: обновление данных пользователя
    elif request.method == 'PUT':
        data = request.json
        
        # Не позволяем удалить последнего администратора
        if user.is_admin and not data.get('is_admin'):
            admin_count = User.query.filter_by(is_admin=True).count()
            if admin_count <= 1:
                return jsonify({'success': False, 'message': 'Невозможно удалить последнего администратора'})
        
        user.first_name = data.get('first_name', user.first_name)
        user.last_name = data.get('last_name', user.last_name)
        user.department = data.get('department', user.department)
        user.position = data.get('position', user.position)
        user.is_admin = data.get('is_admin', user.is_admin)
        
        # Проверяем, если email изменился
        new_email = data.get('email')
        if new_email and new_email != user.email:
            existing_user = User.query.filter_by(email=new_email).first()
            if existing_user:
                return jsonify({'success': False, 'message': 'Этот email уже используется другим пользователем'})
            user.email = new_email
        
        db.session.commit()
        
        # Логируем действие
        log_user_action(f"Обновление пользователя {user.email} (ID: {user.id})")
        
        return jsonify({'success': True})
    
    # DELETE-запрос: удаление пользователя
    elif request.method == 'DELETE':
        # Не позволяем удалить себя
        if user.id == current_user.id:
            return jsonify({'success': False, 'message': 'Вы не можете удалить собственную учетную запись'})
        
        # Не позволяем удалить последнего администратора
        if user.is_admin:
            admin_count = User.query.filter_by(is_admin=True).count()
            if admin_count <= 1:
                return jsonify({'success': False, 'message': 'Невозможно удалить последнего администратора'})
        
        user_email = user.email  # Сохраняем для лога
        db.session.delete(user)
        db.session.commit()
        
        # Логируем действие
        log_user_action(f"Удаление пользователя {user_email}")
        
        return jsonify({'success': True})

@app.route('/api/admin/toggle-admin', methods=['POST'])
@login_required
@admin_required
def toggle_admin():
    data = request.json
    user_id = data.get('user_id')
    is_admin = data.get('is_admin', False)
    
    user = User.query.get_or_404(user_id)
    
    # Не позволяем удалить последнего администратора
    if user.is_admin and not is_admin:
        admin_count = User.query.filter_by(is_admin=True).count()
        if admin_count <= 1:
            return jsonify({'success': False, 'message': 'Невозможно удалить последнего администратора'})
    
    user.is_admin = is_admin
    db.session.commit()
    
    # Логируем действие
    action = "Назначение" if is_admin else "Снятие"
    log_user_action(f"{action} пользователя {user.email} (ID: {user.id}) администратором")
    
    return jsonify({'success': True})

@app.route('/api/admin/feature-access', methods=['POST'])
@login_required
@admin_required
def feature_access():
    data = request.json
    target = data.get('target', '')
    features = data.get('features', {})
    
    if not target or not features:
        return jsonify({'success': False, 'message': 'Неверные данные'})
    
    # Разбираем target вида 'user:1' или 'department:Отдел разработки'
    target_parts = target.split(':', 1)
    if len(target_parts) != 2:
        return jsonify({'success': False, 'message': 'Неверный формат цели'})
    
    target_type, target_id = target_parts
    
    # Удаляем старые настройки доступа
    FeatureAccess.query.filter_by(target_type=target_type, target_id=target_id).delete()
    
    # Добавляем новые настройки доступа
    for feature, access in features.items():
        feature_access = FeatureAccess(
            target_type=target_type,
            target_id=target_id,
            feature=feature,
            access=access
        )
        db.session.add(feature_access)
    
    db.session.commit()
    
    # Логируем действие
    target_desc = f"отдела {target_id}" if target_type == 'department' else f"пользователя (ID: {target_id})"
    log_user_action(f"Обновление прав доступа для {target_desc}")
    
    return jsonify({'success': True})

@app.route('/api/admin/get-feature-access', methods=['GET'])
@login_required
@admin_required
def get_feature_access():
    target = request.args.get('target', '')
    
    if not target:
        return jsonify({'success': False, 'message': 'Не указана цель'})
    
    # Разбираем target вида 'user:1' или 'department:Отдел разработки'
    target_parts = target.split(':', 1)
    if len(target_parts) != 2:
        return jsonify({'success': False, 'message': 'Неверный формат цели'})
    
    target_type, target_id = target_parts
    
    # Получаем настройки доступа из базы данных
    access_records = FeatureAccess.query.filter_by(target_type=target_type, target_id=target_id).all()
    
    # Формируем словарь с настройками доступа
    features = {}
    for record in access_records:
        features[record.feature] = record.access
    
    return jsonify({
        'success': True,
        'target': target,
        'features': features
    })

@app.route('/api/admin/access-summary', methods=['GET'])
@login_required
@admin_required
def access_summary():
    # Получаем все уникальные отделы
    departments = db.session.query(User.department).distinct().all()
    departments = [d[0] for d in departments]
    
    # Получаем всех пользователей
    users = User.query.all()
    user_data = [{'id': u.id, 'name': f"{u.first_name} {u.last_name}", 'email': u.email} for u in users]
    
    # Получаем все записи о доступе
    all_access = FeatureAccess.query.all()
    
    # Формируем словарь настроек доступа
    access_map = {}
    for access in all_access:
        key = f"{access.target_type}:{access.target_id}"
        if key not in access_map:
            access_map[key] = {}
        access_map[key][access.feature] = access.access
    
    # Формируем итоговый результат
    result = {
        'departments': departments,
        'users': user_data,
        'access_map': access_map
    }
    
    return jsonify(result)

@app.route('/api/admin/logs', methods=['POST'])
@login_required
@admin_required
def admin_logs():
    data = request.json
    start_date = data.get('start_date')
    end_date = data.get('end_date')
    
    query = UserLog.query
    
    if start_date:
        start_date = datetime.strptime(start_date, '%Y-%m-%d')
        query = query.filter(UserLog.timestamp >= start_date)
    
    if end_date:
        end_date = datetime.strptime(end_date, '%Y-%m-%d')
        end_date = end_date + timedelta(days=1)  # До конца дня
        query = query.filter(UserLog.timestamp <= end_date)
    
    logs = query.order_by(UserLog.timestamp.desc()).all()
    
    logs_data = []
    for log in logs:
        logs_data.append({
            'id': log.id,
            'user_email': log.user_email,
            'action': log.action,
            'ip_address': log.ip_address or 'Неизвестно',
            'timestamp': log.timestamp.strftime('%d.%m.%Y %H:%M:%S')
        })
    
    return jsonify({'logs': logs_data})

@app.route('/api/load_news', methods=['POST'])
def load_news():
    data = request.json
    start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
    end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date()
    
    results = []
    
    # === arzge.ru ===
    try:
        url_arzge = "https://arzge.ru/news/"
        html_arzge = requests.get(url_arzge, timeout=10).text
        news_items_arzge = parse_arzge(html_arzge, url_arzge, start_date, end_date)
        results.append({"site": "ТАУГАЗ", "news": news_items_arzge})
    except Exception as e:
        results.append({"site": "ТАУГАЗ", "error": f"Ошибка: {e}"})
    
    # === emis-kip.ru ===
    try:
        url_emis = "https://emis-kip.ru/company/sob/news/"
        html_emis = requests.get(url_emis, timeout=10).text
        news_items_emis = parse_emis_kip(html_emis, url_emis, start_date, end_date)
        results.append({"site": "ЭМИС", "news": news_items_emis})
    except Exception as e:
        results.append({"site": "ЭМИС", "error": f"Ошибка: {e}"})

    # === ktkprom.com ===
    try:
        url_ktk = "https://ktkprom.com/novosti-i-sobytija/"
        html_ktk = requests.get(url_ktk, timeout=10).text
        news_items_ktk = parse_ktkprom(html_ktk, start_date, end_date)
        results.append({"site": "КТМПРОМ", "news": news_items_ktk})
    except Exception as e:
        results.append({"site": "КТМПРОМ", "error": f"Ошибка: {e}"})

    # === elmetro.ru ===
    try:
        url_elmetro = "https://www.elmetro.ru/ppecc-tsentp/news/"
        html_elmetro = requests.get(url_elmetro, timeout=10).text
        news_items_elmetro = parse_elmetro(html_elmetro, url_elmetro, start_date, end_date)
        results.append({"site": "ЭЛМЕТРО", "news": news_items_elmetro})
    except Exception as e:
        results.append({"site": "ЭЛМЕТРО", "error": f"Ошибка: {e}"})

    # === vympel.group ===
    try:
        url_vympel = "https://vympel.group/press/"
        html_vympel = requests.get(url_vympel, timeout=10).text
        news_items_vympel = parse_vympel_group(html_vympel, url_vympel, start_date, end_date)
        results.append({"site": "ВЫМПЕЛ", "news": news_items_vympel})
    except Exception as e:
        results.append({"site": "ВЫМПЕЛ", "error": f"Ошибка: {e}"})

    # === packo.ru ===
    try:
        url_packo = "https://packo.ru/novosti"
        html_packo = requests.get(url_packo, timeout=10).text
        news_items_packo = parse_packo(html_packo, url_packo, start_date, end_date)
        results.append({"site": "PACKO", "news": news_items_packo})
    except Exception as e:
        results.append({"site": "PACKO", "error": f"Ошибка: {e}"})

    # === vzljot.ru ===
    try:
        url_vzljot = "https://vzljot.ru/kompaniya/novosti/"
        html_vzljot = requests.get(url_vzljot, timeout=10).text
        news_items_vzljot = parse_vzljot(html_vzljot, url_vzljot, start_date, end_date)
        results.append({"site": "ВЗЛЁТ", "news": news_items_vzljot})
    except Exception as e:
        results.append({"site": "ВЗЛЁТ", "error": f"Ошибка: {e}"})

    return jsonify(results)

def parse_emis_kip(html, base_url, start_date, end_date):
    results = []
    month_map = {
        "января": "January", "февраля": "February", "марта": "March",
        "апреля": "April", "мая": "May", "июня": "June",
        "июля": "July", "августа": "August", "сентября": "September",
        "октября": "October", "ноября": "November", "декабря": "December"
    }

    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select(".news-list__item")

    # Парсим новости на текущей странице
    for item in news_blocks:
        try:
            date_str = item.select_one(".news-list__item-period-date").text.strip()
            for ru, en in month_map.items():
                date_str = date_str.replace(ru, en)
            date = datetime.strptime(date_str, "%d %B %Y").date()
            if start_date <= date <= end_date:
                title_tag = item.select_one(".news-list__item-title a")
                title = title_tag.text.strip()
                href = title_tag['href']
                full_link = urljoin(base_url, href)
                results.append({"date": date.strftime("%d.%m.%Y"), "title": title, "link": full_link})
        except Exception as e:
            print(f"Ошибка при парсинге новости ЭМИС: {e}")
            continue

    return results

def parse_ktkprom(html, start_date, end_date):
    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select(".news_item")
    results = []

    for item in news_blocks:
        try:
            date_str = item.select_one("meta[itemprop='datePublished']")["content"]
            date = datetime.strptime(date_str, "%Y-%m-%d").date()

            if start_date <= date <= end_date:
                title_tag = item.select_one("span[itemprop='headline']")
                title = title_tag.text.strip()

                link_tag = item.select_one("a.news_item_title")
                link = link_tag['href'].strip()

                results.append({"date": date.strftime("%d.%m.%Y"), "title": title, "link": link})
        except Exception as e:
            print(f"Ошибка при парсинге новости КТМПРОМ: {e}")
            continue

    return results

def parse_vzljot(html, base_url, start_date, end_date):
    results = []
    seen_links = set()  # Для отслеживания уникальных ссылок

    # Функция для парсинга одной страницы
    def parse_page(html_content):
        soup = BeautifulSoup(html_content, 'html.parser')
        news_blocks = soup.find_all('div', class_='news-card')

        for block in news_blocks:
            try:
                # Извлечение заголовка
                title_tag = block.find('div', class_='news-card__title')
                if not title_tag:
                    continue
                title = title_tag.text.strip()

                # Извлечение даты
                date_tag = block.find('div', class_='news-card__date')
                if not date_tag:
                    continue
                date_str = date_tag.text.strip()
                date = parse_russian_date(date_str)
                
                # Если дату не удалось распарсить, пропускаем
                if not date:
                    continue

                # Проверка даты
                if not (start_date <= date <= end_date):
                    continue

                # Извлечение ссылки
                link_tag = block.find('a', class_='news-card__link')
                if not link_tag or 'href' not in link_tag.attrs:
                    continue
                relative_link = link_tag['href']
                link = urljoin(base_url, relative_link)

                # Проверка на уникальность ссылки
                if link in seen_links:
                    continue
                seen_links.add(link)

                results.append({
                    "date": date.strftime("%d.%m.%Y"),
                    "title": title,
                    "link": link
                })
            except Exception as e:
                print(f"Ошибка при парсинге блока ВЗЛЁТ: {e}")
                continue

    # Парсинг первой страницы
    parse_page(html)

    # Поиск ссылок на следующие страницы
    soup = BeautifulSoup(html, 'html.parser')
    pagination_items = soup.select(".pagination__item")
    next_page_link = None

    for item in pagination_items:
        if "pagination__item_active" in item.get("class", []):
            # Текущая страница найдена, ищем следующую
            next_sibling = item.find_next_sibling()
            if next_sibling and next_sibling.name == "li":
                next_page_a = next_sibling.find("a")
                if next_page_a and "href" in next_page_a.attrs:
                    next_page_link = next_page_a["href"]
            break

    # Переход по страницам
    while next_page_link:
        try:
            # Получаем HTML-код следующей страницы
            next_page_url = urljoin(base_url, next_page_link)
            response = requests.get(next_page_url, timeout=10)
            if response.status_code != 200:
                break
            next_html = response.text

            # Парсим следующую страницу
            parse_page(next_html)

            # Ищем ссылку на следующую страницу
            soup = BeautifulSoup(next_html, 'html.parser')
            pagination_items = soup.select(".pagination__item")
            next_page_link = None
            for item in pagination_items:
                if "pagination__item_active" in item.get("class", []):
                    next_sibling = item.find_next_sibling()
                    if next_sibling and next_sibling.name == "li":
                        next_page_a = next_sibling.find("a")
                        if next_sibling and next_sibling.name == "li" and next_page_a and "href" in next_page_a.attrs:
                            next_page_link = next_page_a["href"]
                    break
        except Exception as e:
            print(f"Ошибка при переходе на следующую страницу ВЗЛЁТ: {e}")
            break

    return results

def parse_arzge(html, base_url, start_date, end_date):
    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select(".wp-block-columns .wp-block-column.is-layout-flow")
    results = []
    month_map = {
        "января": "January", "февраля": "February", "марта": "March",
        "апреля": "April", "мая": "May", "июня": "June",
        "июля": "July", "августа": "August", "сентября": "September",
        "октября": "October", "ноября": "November", "декабря": "December"
    }

    for block in news_blocks:
        try:
            # Извлечение заголовка
            title_tag = block.select_one("h4.wp-block-heading")
            if not title_tag:
                continue
            title = title_tag.text.strip()

            # Извлечение даты
            date_tag = block.select_one("p.has-text-align-left")
            if not date_tag:
                continue
            date_str = date_tag.text.strip()

            # Очистка строки даты от лишних символов
            date_str = date_str.replace("\u00a0", " ").replace("г.", "").strip()
            for ru, en in month_map.items():
                date_str = date_str.replace(ru, en)

            # Преобразование даты
            try:
                date = datetime.strptime(date_str, "%d %B %Y").date()
            except ValueError:
                print(f"Ошибка при парсинге даты: '{date_str}'")
                continue

            # Проверка даты
            if not (start_date <= date <= end_date):
                continue

            # Извлечение ссылки
            link_tag = block.select_one("a[href]")
            link = link_tag['href'] if link_tag and 'href' in link_tag.attrs else "#"

            # Извлечение изображения
            img_tag = block.select_one("img[src]")
            img_src = img_tag['src'] if img_tag and 'src' in img_tag.attrs else None

            results.append({
                "date": date.strftime("%d.%m.%Y"),
                "title": title,
                "link": link,
                "image": img_src
            })
        except Exception as e:
            print(f"Ошибка при парсинге блока ТАУГАЗ: {e}")
            continue

    return results

def parse_elmetro(html, base_url, start_date, end_date):
    results = []
    month_map = {
        "Января": "January", "Февраля": "February", "Марта": "March",
        "Апреля": "April", "Мая": "May", "Июня": "June",
        "Июля": "July", "Августа": "August", "Сентября": "September",
        "Октября": "October", "Ноября": "November", "Декабря": "December"
    }

    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select(".article__card")

    # Парсим новости на текущей странице
    for item in news_blocks:
        try:
            date_str = item.select_one(".article__card-date").text.strip()
            for ru, en in month_map.items():
                date_str = date_str.replace(ru, en)
            date = datetime.strptime(date_str, "%d %B, %Y").date()
            if start_date <= date <= end_date:
                title_tag = item.select_one(".article__card-title")
                title = title_tag.text.strip()
                href = title_tag['href']
                full_link = urljoin(base_url, href)
                results.append({
                    "date": date.strftime("%d.%m.%Y"),
                    "title": title,
                    "link": full_link
                })
        except Exception as e:
            print(f"Ошибка при парсинге новости ЭЛМЕТРО: {e}")
            continue

    return results

def parse_vympel_group(html, base_url, start_date, end_date):
    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select(".press__item")
    results = []

    month_map = {
        "января": "January", "февраля": "February", "марта": "March",
        "апреля": "April", "мая": "May", "июня": "June",
        "июля": "July", "августа": "August", "сентября": "September",
        "октября": "October", "ноября": "November", "декабря": "December"
    }

    for item in news_blocks:
        try:
            date_str = item.select_one(".press__date").text.strip()
            for ru, en in month_map.items():
                date_str = date_str.replace(ru, en)
            date = datetime.strptime(date_str, "%d %B %Y").date()

            if start_date <= date <= end_date:
                title_tag = item.select_one(".press__title")
                title = title_tag.text.strip()
                href = item.select_one(".press__image")["href"]
                full_link = urljoin(base_url, href)

                results.append({
                    "date": date.strftime("%d.%m.%Y"),
                    "title": title,
                    "link": full_link
                })
        except Exception as e:
            print(f"Ошибка при парсинге новости ВЫМПЕЛ: {e}")
            continue

    return results

def parse_packo(html, base_url, start_date, end_date):
    results = []
    month_map = {
        "января": "January", "февраля": "February", "марта": "March",
        "апреля": "April", "мая": "May", "июня": "June",
        "июля": "July", "августа": "August", "сентября": "September",
        "октября": "October", "ноября": "November", "декабря": "December"
    }

    soup = BeautifulSoup(html, 'html.parser')
    news_blocks = soup.select("a.news_block2")
    
    for block in news_blocks:
        try:
            # Извлечение заголовка
            title_tag = block.select_one(".news_block_title.news_block_title_i")
            if not title_tag:
                continue
            title = title_tag.text.strip()
            
            # Извлечение даты
            date_tag = block.select_one(".news_date")
            if not date_tag:
                continue
            date_str = date_tag.text.strip()
            
            # Замена русских названий месяцев на английские
            for ru_month, en_month in month_map.items():
                date_str = date_str.replace(ru_month, en_month)
            
            # Преобразование строки даты в объект datetime
            date = datetime.strptime(date_str, "%d %B %Y").date()
            
            # Проверка даты
            if not (start_date <= date <= end_date):
                continue
            
            # Извлечение текста новости
            preview_tag = block.select_one(".news_desc.news_desc_i_is_truncated")
            preview = preview_tag.text.strip() if preview_tag else ""
            
            # Извлечение ссылки на полную новость
            link = urljoin(base_url, block['href']) if 'href' in block.attrs else "#"
            
            # Извлечение изображения
            img_tag = block.select_one(".img_cont_block_i img")
            img_src = img_tag['src'] if img_tag and 'src' in img_tag.attrs else None
            if img_src:
                img_src = urljoin(base_url, img_src)  # Преобразуем относительный путь в абсолютный
            
            results.append({
                "date": date.strftime("%d.%m.%Y"),
                "title": title,
                "link": link,
                "preview": preview,
                "image": img_src
            })
        except Exception as e:
            print(f"Ошибка при парсинге блока PACKO: {e}")
            continue

    return results

# Вспомогательная функция для парсинга русских дат
def parse_russian_date(date_str):
    """
    Преобразует дату из строки с русскими названиями месяцев в объект datetime.date.
    Пример входной строки: "12 октября 2023"
    """
    month_map = {
        "января": "January", "февраля": "February", "марта": "March",
        "апреля": "April", "мая": "May", "июня": "June",
        "июля": "July", "августа": "August", "сентября": "September",
        "октября": "October", "ноября": "November", "декабря": "December"
    }
    # Заменяем русские названия месяцев на английские
    for ru_month, en_month in month_map.items():
        date_str = date_str.replace(ru_month, en_month)
    # Преобразуем строку в объект datetime.date
    try:
        return datetime.strptime(date_str, "%d %B %Y").date()
    except ValueError as e:
        print(f"Ошибка при парсинге даты: '{date_str}' - {e}")
        return None

@app.route('/api/generate_qr', methods=['POST'])
def generate_qr():
    data = request.json
    url = data.get('url')
    
    if not url:
        return jsonify({"error": "URL отсутствует"}), 400
    
    # Генерация уникального QR-кода для каждого запроса даже для одинаковых URL
    unique_data = f"{url}_{uuid.uuid4()}"
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(unique_data)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Создаем директорию для QR-кодов, если она не существует
    qr_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'qr_codes')
    os.makedirs(qr_dir, exist_ok=True)
    
    # Генерируем имя файла и сохраняем QR-код
    filename = f"qr_{uuid.uuid4()}.png"
    file_path = os.path.join(qr_dir, filename)
    img.save(file_path)
    
    return jsonify({"qr_code": filename, "file_path": file_path})

@app.route('/qr_codes/<filename>')
def serve_qr_code(filename):
    qr_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'qr_codes')
    return send_file(os.path.join(qr_dir, filename), mimetype='image/png')

@app.route('/api/telegram/parse', methods=['POST'])
def telegram_parse():
    """API-endpoint для парсинга сообщений с канала Telegram"""
    data = request.json
    channel = data.get('channel', '')
    page = data.get('page', 1)
    per_page = data.get('per_page', 20)
    save_to_json = data.get('save_to_json', False)
    # Добавлены новые параметры
    start_date = data.get('start_date')
    end_date = data.get('end_date')
    
    if not channel:
        return jsonify({"error": "Не указан канал Telegram"}), 400
        
    # Удаляем @ из начала, если есть
    if channel.startswith('@'):
        channel = channel[1:]
    
    offset = (page - 1) * per_page
    limit = per_page
    max_messages = offset + limit
    
    # Список для хранения сообщений текущей страницы
    new_messages = []
    
    try:
        # Если запрашивается первая страница, делаем прямой запрос к Telegram
        if page == 1:
            app.logger.info(f"Запрос первой страницы канала {channel}")
            
            # Получаем данные из Telegram
            messages, has_more = parse_telegram_channel(channel, start_date, end_date)
            
            # Сохраняем все полученные сообщения в JSON, если запрошено
            if save_to_json and messages:
                cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
                os.makedirs(cache_dir, exist_ok=True)
                cache_file = os.path.join(cache_dir, f"{channel}.json")
                
                try:
                    # Сохраняем сообщения в файл
                    data_to_save = {
                        'channel': channel,
                        'channel_name': f"@{channel}",
                        'messages': messages,
                        'total_messages': len(messages),
                        'last_update': datetime.now().isoformat()
                    }
                    
                    with open(cache_file, 'w', encoding='utf-8') as f:
                        json.dump(data_to_save, f, ensure_ascii=False)
                        
                    app.logger.info(f"Сохранено {len(messages)} сообщений в кэш для канала {channel}")
                except Exception as e:
                    app.logger.error(f"Ошибка при сохранении кэша: {str(e)}")
            
            # Обрезаем сообщения для текущей страницы
            if len(messages) > offset:
                new_messages = messages[offset:min(max_messages, len(messages))]
            else:
                new_messages = []
                
            return jsonify({
                "messages": new_messages,
                "pagination": {
                    "total_messages": len(messages),
                    "page": page,
                    "per_page": per_page,
                    "has_more": len(messages) > max_messages
                },
                "source": "api"
            })
        
        # Если запрашивается вторая или последующая страница, загружаем из кэша если есть
        else:
            app.logger.info(f"Запрос страницы {page} канала {channel}")
            
            # Проверяем наличие кэша
            cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
            cache_file = os.path.join(cache_dir, f"{channel}.json")
            
            if os.path.exists(cache_file):
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        cache_data = json.load(f)
                    
                    if 'messages' in cache_data and isinstance(cache_data['messages'], list):
                        messages = cache_data['messages']
                        
                        app.logger.info(f"Загружено {len(messages)} сообщений из кэша для канала {channel}")
                        
                        # Обрезаем сообщения для текущей страницы
                        if len(messages) > offset:
                            new_messages = messages[offset:min(max_messages, len(messages))]
                        else:
                            new_messages = []
                        
                        return jsonify({
                            "messages": new_messages,
                            "pagination": {
                                "total_messages": len(messages),
                                "page": page,
                                "per_page": per_page,
                                "has_more": len(messages) > max_messages
                            },
                            "source": "cache"
                        })
                    else:
                        return jsonify({"error": "Некорректный формат кэша"}), 500
                except Exception as e:
                    app.logger.error(f"Ошибка при чтении кэша: {str(e)}")
                    return jsonify({"error": f"Ошибка при чтении кэша: {str(e)}"}), 500
            else:
                # Если кэш не найден, запрашиваем первую страницу и пытаемся загрузить сообщения для пагинации
                app.logger.info(f"Кэш не найден, запрашиваем данные у Telegram")
                
                messages, has_more = parse_telegram_channel(channel, start_date, end_date)
                
                # Если сообщения получены, сохраняем их в кэш
                if messages:
                    os.makedirs(cache_dir, exist_ok=True)
                    try:
                        data_to_save = {
                            'channel': channel,
                            'channel_name': f"@{channel}",
                            'messages': messages,
                            'total_messages': len(messages),
                            'last_update': datetime.now().isoformat()
                        }
                        
                        with open(cache_file, 'w', encoding='utf-8') as f:
                            json.dump(data_to_save, f, ensure_ascii=False)
                            
                        app.logger.info(f"Сохранено {len(messages)} сообщений в кэш для канала {channel}")
                    except Exception as e:
                        app.logger.error(f"Ошибка при сохранении кэша: {str(e)}")
                
                # Обрезаем сообщения для текущей страницы
                if len(messages) > offset:
                    new_messages = messages[offset:min(max_messages, len(messages))]
                else:
                    new_messages = []
                
                return jsonify({
                    "messages": new_messages,
                    "pagination": {
                        "total_messages": len(messages),
                        "page": page,
                        "per_page": per_page,
                        "has_more": len(messages) > max_messages
                    },
                    "source": "api_fallback"
                })
                
    except Exception as e:
        app.logger.error(f"Ошибка при парсинге канала {channel}: {str(e)}")
        return jsonify({
            "error": f"Ошибка при загрузке данных: {str(e)}"
        }), 500

@app.route('/api/telegram/favorites', methods=['GET'])
@login_required
def get_favorites():
    # Получаем избранные каналы текущего пользователя из базы данных
    favorites = UserFavorite.query.filter_by(user_id=current_user.id).all()
    channels = [favorite.channel for favorite in favorites]
    return jsonify(channels)

@app.route('/api/telegram/favorites', methods=['POST'])
@login_required
def add_favorite():
    data = request.json
    channel = data.get('channel', '')
    
    if not channel:
        return jsonify({"error": "Не указан канал Telegram"}), 400
        
    # Удаляем @ из начала, если есть
    if channel.startswith('@'):
        channel = channel[1:]
    
    # Проверяем, есть ли уже такой канал в избранном у пользователя
    existing_favorite = UserFavorite.query.filter_by(
        user_id=current_user.id, 
        channel=channel
    ).first()
    
    if not existing_favorite:
        # Добавляем новый канал в избранное
        new_favorite = UserFavorite(user_id=current_user.id, channel=channel)
        db.session.add(new_favorite)
        db.session.commit()
    
    # Получаем обновленный список избранных
    favorites = UserFavorite.query.filter_by(user_id=current_user.id).all()
    channels = [favorite.channel for favorite in favorites]
    
    return jsonify({"status": "success", "favorites": channels})

# Новый маршрут для удаления канала из избранного
@app.route('/api/telegram/favorites/<channel>', methods=['DELETE'])
@login_required
def remove_favorite(channel):
    # Находим запись в избранном
    favorite = UserFavorite.query.filter_by(
        user_id=current_user.id, 
        channel=channel
    ).first()
    
    if favorite:
        db.session.delete(favorite)
        db.session.commit()
        return jsonify({"status": "success", "message": "Канал удален из избранного"})
    else:
        return jsonify({"error": "Канал не найден в избранном"}), 404

@app.route('/api/export_news_excel', methods=['POST'])
def export_news_excel():
    try:
        # Получаем данные из запроса
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        export_type = data.get('type', 'holdings')  # По умолчанию - холдинги
        
        # Формируем имя файла
        now = datetime.now()
        type_prefix = "Конкуренты" if export_type == 'competitors' else "Холдинги"
        file_name = f"Новости_{type_prefix}_{now.strftime('%d-%m-%Y_%H-%M-%S')}.xlsx"
        
        print(f"Начинаем экспорт в Excel. Тип: {export_type}, Период: {start_date} - {end_date}")
        
        # Если это экспорт новостей конкурентов, загружаем данные
        if export_type == 'competitors':
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            
            results = []
            
            # === arzge.ru ===
            try:
                url_arzge = "https://arzge.ru/news/"
                html_arzge = requests.get(url_arzge, timeout=10).text
                news_items_arzge = parse_arzge(html_arzge, url_arzge, start_date_obj, end_date_obj)
                results.append({"site": "ТАУГАЗ", "news": news_items_arzge})
            except Exception as e:
                results.append({"site": "ТАУГАЗ", "error": f"Ошибка: {e}"})
            
            # === emis-kip.ru ===
            try:
                url_emis = "https://emis-kip.ru/company/sob/news/"
                html_emis = requests.get(url_emis, timeout=10).text
                news_items_emis = parse_emis_kip(html_emis, url_emis, start_date_obj, end_date_obj)
                results.append({"site": "ЭМИС", "news": news_items_emis})
            except Exception as e:
                results.append({"site": "ЭМИС", "error": f"Ошибка: {e}"})

            # === ktkprom.com ===
            try:
                url_ktk = "https://ktkprom.com/novosti-i-sobytija/"
                html_ktk = requests.get(url_ktk, timeout=10).text
                news_items_ktk = parse_ktkprom(html_ktk, start_date_obj, end_date_obj)
                results.append({"site": "КТМПРОМ", "news": news_items_ktk})
            except Exception as e:
                results.append({"site": "КТМПРОМ", "error": f"Ошибка: {e}"})

            # === elmetro.ru ===
            try:
                url_elmetro = "https://www.elmetro.ru/ppecc-tsentp/news/"
                html_elmetro = requests.get(url_elmetro, timeout=10).text
                news_items_elmetro = parse_elmetro(html_elmetro, url_elmetro, start_date_obj, end_date_obj)
                results.append({"site": "ЭЛМЕТРО", "news": news_items_elmetro})
            except Exception as e:
                results.append({"site": "ЭЛМЕТРО", "error": f"Ошибка: {e}"})

            # === vympel.group ===
            try:
                url_vympel = "https://vympel.group/press/"
                html_vympel = requests.get(url_vympel, timeout=10).text
                news_items_vympel = parse_vympel_group(html_vympel, url_vympel, start_date_obj, end_date_obj)
                results.append({"site": "ВЫМПЕЛ", "news": news_items_vympel})
            except Exception as e:
                results.append({"site": "ВЫМПЕЛ", "error": f"Ошибка: {e}"})

            # === packo.ru ===
            try:
                url_packo = "https://packo.ru/novosti"
                html_packo = requests.get(url_packo, timeout=10).text
                news_items_packo = parse_packo(html_packo, url_packo, start_date_obj, end_date_obj)
                results.append({"site": "PACKO", "news": news_items_packo})
            except Exception as e:
                results.append({"site": "PACKO", "error": f"Ошибка: {e}"})

            # === vzljot.ru ===
            try:
                url_vzljot = "https://vzljot.ru/kompaniya/novosti/"
                html_vzljot = requests.get(url_vzljot, timeout=10).text
                news_items_vzljot = parse_vzljot(html_vzljot, url_vzljot, start_date_obj, end_date_obj)
                results.append({"site": "ВЗЛЁТ", "news": news_items_vzljot})
            except Exception as e:
                results.append({"site": "ВЗЛЁТ", "error": f"Ошибка: {e}"})
                
            data = results
        
        # Создаем временный файл для Excel
        tmp_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
        print(f"Создаем Excel файл: {tmp_path}")
        
        # Создаем Excel файл
        import xlsxwriter
        
        # Создаем книгу Excel
        workbook = xlsxwriter.Workbook(tmp_path)
        
        # Создаем форматы для заголовков и ячеек
        header_format = workbook.add_format({
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'bg_color': '#D9E1F2'
        })
        
        cell_format = workbook.add_format({
            'valign': 'vcenter',
            'border': 1,
            'text_wrap': True
        })
        
        date_format = workbook.add_format({
            'valign': 'vcenter',
            'border': 1,
            'num_format': 'dd.mm.yyyy'
        })
        
        # Для каждого сайта создаем отдельный лист
        for site in data:
            if not site.get('error') and site.get('news') and len(site['news']) > 0:
                # Имя листа (ограничено 31 символом)
                sheet_name = site['site'][:31]
                
                # Создаем лист
                worksheet = workbook.add_worksheet(sheet_name)
                
                # Устанавливаем ширину колонок
                worksheet.set_column('A:A', 15)  # Дата
                worksheet.set_column('B:B', 60)  # Название новости
                worksheet.set_column('C:C', 50)  # Ссылка
                
                # Добавляем заголовки
                worksheet.write(0, 0, 'Дата', header_format)
                worksheet.write(0, 1, 'Название новости', header_format)
                worksheet.write(0, 2, 'Ссылка', header_format)
                
                # Добавляем данные
                row = 1
                for item in site['news']:
                    worksheet.write(row, 0, item['date'], date_format)
                    worksheet.write(row, 1, item['title'], cell_format)
                    worksheet.write(row, 2, item['link'], cell_format)
                    row += 1
        
        # Закрываем книгу
        workbook.close()
        print(f"Excel файл успешно создан: {tmp_path}")
        
        # Отправляем файл клиенту
        response = send_file(
            tmp_path,
            as_attachment=True,
            download_name=file_name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
        # Добавляем дополнительные заголовки для обеспечения загрузки
        response.headers["Content-Disposition"] = f"attachment; filename={file_name}"
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Cache-Control"] = "no-cache"
        
        print(f"Файл отправлен клиенту: {file_name}")
        return response
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Ошибка при экспорте в Excel: {str(e)}")
        print(f"Детали ошибки: {error_details}")
        return jsonify({"error": str(e), "details": error_details}), 500

@app.route('/api/ai_chat', methods=['POST'])
def ai_chat():
    data = request.json
    model = data.get('model')
    message = data.get('message')
    api_key = data.get('api_key')
    
    if not all([model, message, api_key]):
        return jsonify({"error": "Не все необходимые данные предоставлены"}), 400
    
    # Здесь будет код для обращения к различным API моделям
    # В качестве заглушки вернем эхо-ответ
    response = {
        "model": model,
        "response": f"Ответ от модели {model}: {message}"
    }
    
    return jsonify(response)

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "Файл не найден"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Имя файла отсутствует"}), 400
    
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)
    
    return jsonify({"success": True, "filename": file.filename, "path": file_path})

@app.route('/api/analyze', methods=['POST'])
def analyze_excel():
    print("Получен запрос на анализ Excel файла")
    
    if 'file' not in request.files:
        print("Ошибка: файл не найден в запросе")
        print(f"Заголовки запроса: {request.headers}")
        print(f"Ключи в request.files: {list(request.files.keys())}")
        return jsonify({"error": "Файл не найден в запросе"}), 400
    
    file = request.files['file']
    if file.filename == '':
        print("Ошибка: имя файла отсутствует")
        return jsonify({"error": "Имя файла отсутствует"}), 400
    
    print(f"Получен файл: {file.filename}")
    
    try:
        # Получаем список листов Excel
        print("Получение списка листов Excel...")
        excel_file = pd.ExcelFile(file, engine='openpyxl')
        sheet_names = excel_file.sheet_names
        print(f"Найдены листы: {sheet_names}")
        
        # Если указан конкретный лист, используем его
        sheet_name = request.form.get('sheet_name')
        if sheet_name and sheet_name in sheet_names:
            print(f"Используем указанный лист: {sheet_name}")
            df = pd.read_excel(excel_file, sheet_name=sheet_name, engine='openpyxl')
        else:
            # Используем первый лист по умолчанию
            print(f"Используем первый лист: {sheet_names[0]}")
            df = pd.read_excel(excel_file, sheet_name=sheet_names[0], engine='openpyxl')
        
        print(f"Файл успешно прочитан. Размер DataFrame: {df.shape}")
        print(f"Столбцы: {df.columns.tolist()}")
        
        # Заменяем NaN, inf и -inf на None для корректной JSON-сериализации
        df = df.replace([np.nan, np.inf, -np.inf], [None, None, None])
        
        # Преобразование DataFrame в формат JSON для передачи клиенту
        columns = df.columns.tolist()
        data = df.values.tolist()
        
        # Добавляем имена столбцов в начало списка данных
        data.insert(0, columns)
        
        # Проверяем на наличие проблемных значений в данных
        for i, row in enumerate(data):
            for j, value in enumerate(row):
                # Заменяем проблемные значения на None
                if isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
                    data[i][j] = None
        
        return jsonify({
            "success": True,
            "data": data,
            "sheets": sheet_names,
            "current_sheet": sheet_name if sheet_name and sheet_name in sheet_names else sheet_names[0]
        })
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Ошибка при обработке файла: {str(e)}")
        print(f"Traceback: {error_traceback}")
        return jsonify({"error": f"Ошибка при обработке файла: {str(e)}"}), 500

@app.route('/api/analyze_columns', methods=['POST'])
def analyze_columns():
    print("Начинаем анализ столбцов данных")
    try:
        data = request.json
        if not data or 'data' not in data or 'selectedColumns' not in data:
            return jsonify({"error": "Неверные данные запроса"}), 400
        
        raw_data = data['data']
        selected_columns = data['selectedColumns']
        
        # Проверяем, что данные не пустые
        if not raw_data or len(raw_data) < 2:  # Должен быть хотя бы заголовок и одна строка данных
            return jsonify({"error": "Недостаточно данных для анализа"}), 400
        
        print(f"Получены данные: {len(raw_data)} строк, выбрано {len(selected_columns)} столбцов")
        
        # Получаем заголовки всех столбцов
        all_headers = raw_data[0]
        
        # 1. Создаем массив с абсолютными значениями только выбранных столбцов
        absolute_data = []
        selected_headers = [all_headers[idx] for idx in selected_columns]
        absolute_data.append(selected_headers)  # Добавляем заголовки
        
        # Добавляем строки данных
        for i in range(1, len(raw_data)):
            row = raw_data[i]
            if not isinstance(row, list):
                row = list(row)  # Преобразуем в список, если это не список
            
            # Добавляем только выбранные столбцы
            selected_row = [row[idx] for idx in selected_columns]
            absolute_data.append(selected_row)
        
        print(f"Создан массив абсолютных данных: {len(absolute_data)} строк")
        
        # 2. Находим числовые столбцы и годовые столбцы для анализа процентов
        numeric_columns_idx = []
        year_columns_idx = []
        text_columns_idx = []
        
        for col_idx, col_name in enumerate(selected_headers):
            # Проверяем, похож ли заголовок на год (4 цифры начинающиеся с 20)
            is_year_column = False
            
            # Преобразуем заголовок в строку для надежности
            col_name_str = str(col_name)
            
            # Проверка на годовой столбец (более гибкая)
            if (
                (isinstance(col_name, (int, float)) and 2000 <= col_name <= 2100) or 
                (isinstance(col_name_str, str) and col_name_str.isdigit() and len(col_name_str) == 4 and col_name_str.startswith('20'))
            ):
                is_year_column = True
                year_columns_idx.append(col_idx)
                print(f"Обнаружен годовой столбец: {col_name_str} (индекс {col_idx})")
            
            # Проверяем первые несколько строк, чтобы определить, является ли столбец числовым
            is_numeric = True
            for row_idx in range(1, min(len(absolute_data), 5)):
                val = absolute_data[row_idx][col_idx]
                if val is not None:
                    if not isinstance(val, (int, float)):
                        # Пробуем преобразовать в число
                        try:
                            float(val)
                        except (ValueError, TypeError):
                            is_numeric = False
                            break
            
            if is_numeric and not is_year_column:
                numeric_columns_idx.append(col_idx)
            elif not is_numeric and not is_year_column:
                text_columns_idx.append(col_idx)
        
        print(f"Найдено {len(numeric_columns_idx)} числовых столбцов, {len(year_columns_idx)} годовых столбцов и {len(text_columns_idx)} текстовых столбцов")
        
        # Добавляем годовые столбцы к числовым для анализа
        all_numeric_columns = numeric_columns_idx + year_columns_idx
        
        # 3. Преобразуем строковые числовые значения в фактические числа
        for row_idx in range(1, len(absolute_data)):
            for col_idx in all_numeric_columns:
                val = absolute_data[row_idx][col_idx]
                if val is not None and not isinstance(val, (int, float)):
                    try:
                        absolute_data[row_idx][col_idx] = float(val)
                    except (ValueError, TypeError):
                        pass  # Оставляем как есть, если не получилось преобразовать
        
        # 4. Копируем массив данных для процентного соотношения
        percent_data = []
        percent_headers = list(selected_headers)  # Новый список заголовков
        
        # Создаем новый список процентных заголовков
        for col_idx in all_numeric_columns:
            percent_headers.append(f"{selected_headers[col_idx]} (% от рынка)")
        
        percent_data.append(percent_headers)
        
        # Копируем строки данных
        for row_idx in range(1, len(absolute_data)):
            row = list(absolute_data[row_idx])  # Копируем строку
            for col_idx in all_numeric_columns:
                # Добавляем пустые значения для процентных столбцов, заполним позже
                row.append(None)
            percent_data.append(row)
        
        # 5. Вычисляем суммы по каждому числовому столбцу
        column_sums = {}
        for col_idx in all_numeric_columns:
            total = 0
            for row_idx in range(1, len(absolute_data)):
                val = absolute_data[row_idx][col_idx]
                if isinstance(val, (int, float)) and not math.isnan(val) and not math.isinf(val):
                    total += val
            column_sums[col_idx] = total
        
        # 6. Вычисляем процентные соотношения
        for col_idx in all_numeric_columns:
            total = column_sums[col_idx]
            if total > 0:
                percent_col_idx = len(selected_headers) + all_numeric_columns.index(col_idx)
                
                for row_idx in range(1, len(percent_data)):
                    val = absolute_data[row_idx][col_idx]
                    if isinstance(val, (int, float)) and not math.isnan(val) and not math.isinf(val):
                        percent = (val / total) * 100
                        percent_data[row_idx][percent_col_idx] = round(percent, 2)
        
        # 7. Проверяем и заменяем NaN, inf, -inf на None
        for data_list in [absolute_data, percent_data]:
            for row_idx in range(len(data_list)):
                for col_idx in range(len(data_list[row_idx])):
                    val = data_list[row_idx][col_idx]
                    if isinstance(val, float) and (math.isnan(val) or math.isinf(val)):
                        data_list[row_idx][col_idx] = None
        
        # 8. Подготавливаем метаданные для клиента
        metadata = {
            "numeric_columns": numeric_columns_idx,
            "year_columns": year_columns_idx,
            "text_columns": text_columns_idx,
            "available_years": [selected_headers[idx] for idx in year_columns_idx]
        }
        
        # Проверяем и преобразуем годы в строки для большей надежности
        metadata["available_years"] = [str(year) for year in metadata["available_years"]]
        
        print(f"Подготовлены данные: абсолютные {len(absolute_data)}x{len(absolute_data[0])}, процентные {len(percent_data)}x{len(percent_data[0])}")
        print(f"Метаданные: {metadata}")
        
        return jsonify({
            "success": True,
            "absolute_data": absolute_data,
            "percent_data": percent_data,
            "metadata": metadata
        })
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Ошибка при анализе данных: {str(e)}")
        print(f"Traceback: {error_traceback}")
        return jsonify({"error": f"Ошибка при анализе данных: {str(e)}"}), 500

@app.route('/api/test-upload', methods=['POST'])
def test_upload():
    """
    Тестовый endpoint для отладки проблем с загрузкой файлов
    """
    print("Вызов test-upload endpoint")
    response_data = {"files": {}, "form": {}, "headers": {}}
    
    # Информация о запросе
    print(f"Content-Type: {request.content_type}")
    print(f"Данные формы: {request.form}")
    print(f"Файлы: {list(request.files.keys())}")
    
    # Сохраняем информацию о заголовках
    for header, value in request.headers:
        response_data["headers"][header] = value
        print(f"Заголовок: {header} = {value}")
    
    # Сохраняем информацию о полях формы
    for key, value in request.form.items():
        response_data["form"][key] = value
    
    # Сохраняем информацию о файлах
    for file_key in request.files:
        file = request.files[file_key]
        response_data["files"][file_key] = {
            "filename": file.filename,
            "content_type": file.content_type,
            "size": "stream"  # Невозможно определить размер потока без чтения
        }
        print(f"Файл: {file_key} = {file.filename} (тип: {file.content_type})")
    
    return jsonify(response_data)

# Инициализация базы данных при первом запуске
# @app.before_first_request
# def create_tables():
#     db.create_all()

# Создаем все таблицы при запуске приложения
with app.app_context():
    db.create_all()
    
    # Инициализируем кэш новостей, если он пуст
    if NewsCache.query.count() == 0:
        try:
            print("Инициализация кэша новостей...")
            # Отключаем автоматическую загрузку новостей при старте
            # news_items = fetch_news(use_cache=False)
            # print(f"Загружено {len(news_items)} новостей.")
            print("Инициализация кэша новостей отключена.")
        except Exception as e:
            print(f"Ошибка при инициализации кэша новостей: {str(e)}")

# Маршруты для аутентификации и управления профилем
@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = RegistrationForm()
    if form.validate_on_submit():
        hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        
        # Проверяем, есть ли уже пользователи в системе
        users_count = User.query.count()
        
        user = User(
            email=form.email.data,
            password=hashed_password,
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            department=form.department.data,
            position=form.position.data,
            # Если это первый пользователь - делаем его администратором
            is_admin=(users_count == 0)
        )
        db.session.add(user)
        db.session.commit()
        
        # Если это первый пользователь, логируем создание администратора
        if users_count == 0:
            flash('Вы зарегистрированы как первый пользователь и назначены администратором системы!', 'success')
        else:
            flash('Ваш аккаунт был создан! Теперь вы можете войти в систему.', 'success')
        
        return redirect(url_for('login'))
    
    return render_template('register.html', title='Регистрация', form=form)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            # Сохраняем IP адрес и время входа
            user.last_ip = request.remote_addr
            user.last_login = datetime.utcnow()
            db.session.commit()
            
            # Использование remember=True для долгосрочной сессии
            login_user(user, remember=form.remember.data)
            
            next_page = request.args.get('next')
            flash('Вы успешно вошли в систему!', 'success')
            return redirect(next_page) if next_page else redirect(url_for('index'))
        else:
            flash('Ошибка авторизации. Пожалуйста, проверьте email и пароль.', 'danger')
    
    return render_template('login.html', title='Вход', form=form)

@app.route('/logout')
def logout():
    logout_user()
    flash('Вы вышли из системы.', 'info')
    return redirect(url_for('index'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    form = UpdateProfileForm()
    
    if request.method == 'POST':
        # Проверяем, был ли запрос на загрузку фото
        if request.files and 'picture' in request.files and request.files['picture'].filename:
            print("Загружается фото профиля")
            try:
                picture_file = save_picture(request.files['picture'])
                current_user.profile_image = picture_file
                db.session.commit()
                flash('Фото профиля успешно обновлено!', 'success')
                print(f"Фото профиля обновлено: {picture_file}")
                return redirect(url_for('profile'))
            except Exception as e:
                flash(f'Ошибка при сохранении фото: {str(e)}', 'danger')
                print(f"Ошибка при сохранении фото: {str(e)}")
                return redirect(url_for('profile'))
    
    if form.validate_on_submit():
        print("Форма прошла валидацию! Обновляем профиль пользователя")
        print(f"Email: {form.email.data}")
        print(f"Имя: {form.first_name.data}")
        print(f"Фамилия: {form.last_name.data}")
        print(f"Отдел: {form.department.data}")
        print(f"Должность: {form.position.data}")
        
        if form.picture.data:
            picture_file = save_picture(form.picture.data)
            current_user.profile_image = picture_file
            print(f"Фото профиля обновлено: {picture_file}")
        
        # Сохраняем старые значения для сравнения
        old_email = current_user.email
        old_first_name = current_user.first_name
        old_last_name = current_user.last_name
        old_department = current_user.department
        old_position = current_user.position
        
        # Обновляем данные пользователя
        current_user.email = form.email.data
        current_user.first_name = form.first_name.data
        current_user.last_name = form.last_name.data
        current_user.department = form.department.data
        current_user.position = form.position.data
        
        # Выводим изменения
        if old_email != current_user.email:
            print(f"Email изменен: {old_email} -> {current_user.email}")
        if old_first_name != current_user.first_name:
            print(f"Имя изменено: {old_first_name} -> {current_user.first_name}")
        if old_last_name != current_user.last_name:
            print(f"Фамилия изменена: {old_last_name} -> {current_user.last_name}")
        if old_department != current_user.department:
            print(f"Отдел изменен: {old_department} -> {current_user.department}")
        if old_position != current_user.position:
            print(f"Должность изменена: {old_position} -> {current_user.position}")
        
        db.session.commit()
        print("Изменения сохранены в базе данных")
        
        flash('Ваш профиль успешно обновлен!', 'success')
        return redirect(url_for('profile'))
    elif request.method == 'POST':
        print("Форма не прошла валидацию")
        print(f"Ошибки: {form.errors}")
    elif request.method == 'GET':
        form.email.data = current_user.email
        form.first_name.data = current_user.first_name
        form.last_name.data = current_user.last_name
        form.department.data = current_user.department
        form.position.data = current_user.position
        print(f"Загрузка данных пользователя: {current_user.email}, {current_user.first_name}, {current_user.last_name}, {current_user.department}, {current_user.position}")
    
    password_form = UpdatePasswordForm()
    
    return render_template('profile.html', title='Профиль', form=form, 
                          password_form=password_form)

@app.route('/update_password', methods=['POST'])
@login_required
def update_password():
    form = UpdatePasswordForm()
    
    if form.validate_on_submit():
        if bcrypt.check_password_hash(current_user.password, form.current_password.data):
            hashed_password = bcrypt.generate_password_hash(form.new_password.data).decode('utf-8')
            current_user.password = hashed_password
            db.session.commit()
            flash('Ваш пароль был успешно обновлен!', 'success')
        else:
            flash('Текущий пароль неверный. Попробуйте еще раз.', 'danger')
    
    return redirect(url_for('profile'))

# Маршрут для получения профильного изображения
@app.route('/uploads/profile_pics/<filename>')
def profile_pics(filename):
    # Если это стандартный аватар, берем его из static/img
    if filename == 'avatar.png':
        return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'img', 'avatar.png'))
    # Иначе ищем в директории профилей
    return send_file(os.path.join(PROFILE_PICS_DIR, filename))

# Дополнительный маршрут для совместимости со static
@app.route('/static/uploads/profile_pics/<filename>')
def static_profile_pics(filename):
    # Если это стандартный аватар, берем его из static/img
    if filename == 'avatar.png':
        return send_file(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'img', 'avatar.png'))
    # Иначе ищем в директории профилей
    return send_file(os.path.join(PROFILE_PICS_DIR, filename))

@app.route('/api/refresh_news', methods=['GET'])
def refresh_news_api():
    """API-endpoint для асинхронного обновления новостей"""
    try:
        # Получаем актуальные новости с сайта
        fresh_news = fetch_news(use_cache=False)
        
        # Возвращаем новости в формате JSON
        return jsonify({
            'success': True,
            'news': fresh_news
        })
    except Exception as e:
        import traceback
        print(f"Ошибка при обновлении новостей: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        })

def fetch_news(use_cache=True, max_items=10):
    """
    Получает новости с сайта или из кэша.
    
    :param use_cache: Если True, сначала проверяет кэш. Если False, принудительно обновляет новости с сайта.
    :param max_items: Максимальное количество новостей для возврата
    :return: Список новостей [{'title': ..., 'link': ..., 'date': ...}, ...]
    """
    # Если use_cache=True, проверяем наличие актуальных новостей в кэше
    if use_cache:
        # Проверяем, есть ли в кэше свежие новости (не старше 1 часа)
        cache_expiry = datetime.utcnow() - timedelta(hours=1)
        cached_news = NewsCache.query.filter(NewsCache.timestamp >= cache_expiry).order_by(NewsCache.id.desc()).limit(max_items).all()
        
        # Если в кэше есть новости, возвращаем их
        if cached_news:
            print(f"Используем {len(cached_news)} новостей из кэша")
            return NewsCache.to_json(cached_news)
    
    # Если кэш пуст или устарел, или use_cache=False, загружаем новости с сайта
    url = "https://turbo-don.ru"
    try:
        print("Загружаем свежие новости с сайта")
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        news_items = []
        for item in soup.select('.page__lenta-item'):
            title_tag = item.select_one('.page__lenta-item_content')
            date_tag = item.select_one('.page__lenta-item_date')
            link = title_tag['href']
            title = title_tag.text.strip()
            date = date_tag.text.strip() if date_tag else "Дата не указана"
            news_items.append({'title': title, 'link': link, 'date': date})

        # Сохраняем загруженные новости в кэш, если это не асинхронное обновление
        if not use_cache:
            try:
                # Удаляем все старые записи из кэша
                NewsCache.query.delete()
                db.session.commit()
                
                # Добавляем новые записи в кэш
                for item in news_items:
                    cache_item = NewsCache(
                        title=item['title'],
                        link=item['link'],
                        date=item['date']
                    )
                    db.session.add(cache_item)
                db.session.commit()
                print(f"Сохранили {len(news_items)} новостей в кэш")
            except Exception as e:
                db.session.rollback()
                print(f"Ошибка при сохранении новостей в кэш: {str(e)}")
        
        return news_items[:max_items]
    except requests.exceptions.RequestException as e:
        print(f"Error fetching news: {e}")
        
        # В случае ошибки, если кэш отключен, проверяем, есть ли хоть какие-то новости в кэше
        if not use_cache:
            cached_news = NewsCache.query.order_by(NewsCache.id.desc()).limit(max_items).all()
            if cached_news:
                print(f"Не удалось загрузить новости с сайта, используем {len(cached_news)} новостей из кэша")
                return NewsCache.to_json(cached_news)
        
        return []

@app.route('/api/telegram/get_cache', methods=['GET'])
def get_telegram_cache():
    """API-endpoint для получения данных из JSON-кэша"""
    channel = request.args.get('channel', '')
    
    if not channel:
        return jsonify({"error": "Не указан канал Telegram", "success": False}), 400
        
    if channel.startswith('@'):
        channel = channel[1:]
    
    try:
        # Путь к файлу кэша для данного канала
        cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
        
        # Проверяем существование директории кэша и создаем её при необходимости
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)
            app.logger.info(f"Создана директория кэша: {cache_dir}")
        
        cache_file = os.path.join(cache_dir, f"{channel}.json")
        
        app.logger.info(f"Проверка кэша для канала {channel} в файле {cache_file}")
        
        # Проверяем существование файла кэша
        if not os.path.exists(cache_file):
            app.logger.info(f"Кэш для канала {channel} не найден")
            return jsonify({"success": False, "messages": [], "error": "Кэш не найден"})
        
        # Проверяем, что файл не пустой и является корректным JSON
        if os.path.getsize(cache_file) == 0:
            app.logger.warning(f"Файл кэша для канала {channel} пуст")
            return jsonify({"success": False, "messages": [], "error": "Файл кэша пуст"})
        
        try:
            # Загружаем данные из файла кэша
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            messages = cache_data.get('messages', [])
            total_messages = cache_data.get('total_messages', len(messages))
            last_update = cache_data.get('last_update')
            
            app.logger.info(f"Загружено {len(messages)} сообщений из кэша для канала {channel}")
            
            # Проверяем, есть ли сообщения в кэше
            if not messages:
                app.logger.warning(f"Кэш для канала {channel} не содержит сообщений")
                return jsonify({"success": True, "messages": [], "error": "Кэш пуст"})
            
            # Проверяем актуальность кэша (не старше 24 часов)
            cache_is_stale = False
            if last_update:
                try:
                    last_update_time = datetime.fromisoformat(last_update.replace('Z', '+00:00'))
                    cache_age = datetime.now(last_update_time.tzinfo) - last_update_time
                    
                    # Если кэш устарел, помечаем его как устаревший, но всё равно возвращаем данные
                    if cache_age.total_seconds() >= 86400:  # 24 часа
                        app.logger.info(f"Кэш для канала {channel} устарел (возраст: {cache_age.total_seconds() / 3600:.1f} часов)")
                        cache_is_stale = True
                except Exception as e:
                    app.logger.error(f"Ошибка при проверке даты обновления кэша: {str(e)}")
                    # В случае ошибки при парсинге даты, считаем кэш устаревшим
                    cache_is_stale = True
            else:
                # Если нет даты обновления, считаем кэш устаревшим
                cache_is_stale = True
            
            return jsonify({
                "success": True,
                "messages": messages,
                "total_messages": total_messages,
                "last_update": last_update,
                "cache_is_stale": cache_is_stale
            })
            
        except json.JSONDecodeError as e:
            app.logger.error(f"Ошибка при декодировании JSON из кэша: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"Ошибка формата JSON в кэше: {str(e)}",
                "messages": []
            }), 500
            
    except Exception as e:
        app.logger.error(f"Ошибка при получении данных из кэша: {str(e)}", exc_info=True)
        return jsonify({
            "success": False,
            "error": f"Ошибка при получении данных из кэша: {str(e)}",
            "messages": []
        }), 500

@app.route('/api/telegram/load_all', methods=['POST'])
def telegram_load_all():
    """API-endpoint для загрузки ВСЕХ сообщений канала в JSON файл"""
    data = request.json
    channel = data.get('channel', '')
    max_pages = data.get('max_pages', None)  # Опциональное ограничение на количество страниц
    
    if not channel:
        return jsonify({"error": "Не указан канал Telegram"}), 400
        
    # Удаляем @ из начала, если есть
    if channel.startswith('@'):
        channel = channel[1:]
    
    try:
        # Создаем директорию для кэша, если её нет
        cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)
            
        cache_file = os.path.join(cache_dir, f"{channel}.json")
        
        # Промежуточный файл для сохранения прогресса на случай ошибки
        temp_cache_file = os.path.join(cache_dir, f"{channel}_temp.json")
        
        # Устанавливаем флаги для отслеживания прогресса
        messages_loaded = 0
        pages_loaded = 0
        partial_success = False
        
        # Загружаем существующие сообщения из кэша, если он существует
        existing_messages = []
        if os.path.exists(cache_file):
            try:
                print(f"Обнаружен существующий кэш для канала @{channel}, загружаем его...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    existing_data = json.load(f)
                
                if 'messages' in existing_data and len(existing_data['messages']) > 0:
                    existing_messages = existing_data.get('messages', [])
                    print(f"Загружено {len(existing_messages)} существующих сообщений из кэша")
            except Exception as e:
                print(f"Ошибка при загрузке существующего кэша: {str(e)}")
        
        # Создаем словарь для отслеживания уникальных ID сообщений
        message_ids = {msg.get('id'): True for msg in existing_messages if msg.get('id')}
        print(f"Найдено {len(message_ids)} уникальных ID сообщений в существующем кэше")
        
        # Получаем первую страницу сообщений
        print(f"Загрузка сообщений канала @{channel} (макс. страниц: {max_pages or 'без ограничений'})")
        
        # Вызываем parse_telegram_channel и корректно обрабатываем возвращаемые значения
        result = parse_telegram_channel(channel)
        if len(result) == 3:
            initial_messages, has_more, captcha_html = result
        else:
            initial_messages, has_more = result
            captcha_html = None
        
        # Проверяем наличие CAPTCHA
        if captcha_html:
            print(f"Обнаружена CAPTCHA для канала @{channel}")
            return jsonify({
                "success": False,
                "error": "Обнаружена CAPTCHA от Telegram. Необходимо пройти проверку.",
                "captcha_detected": True,
                "channel": channel
            }), 403  # Используем код 403 для индикации необходимости проверки
        
        # Увеличиваем счетчик загруженных страниц
        pages_loaded += 1
        
        # Если первая страница загружена успешно, начинаем обработку
        if initial_messages and len(initial_messages) > 0:
            # Добавляем новые сообщения к существующим, исключая дубликаты
            new_messages_count = 0
            for msg in initial_messages:
                if msg.get('id') and not message_ids.get(msg.get('id')):
                    existing_messages.append(msg)
                    message_ids[msg.get('id')] = True
                    new_messages_count += 1
            
            print(f"Получено {len(initial_messages)} сообщений на первой странице, из них {new_messages_count} новых")
            messages_loaded = len(existing_messages)
            
            # Сохраняем промежуточный результат
            temp_data = {
                'channel': channel,
                'channel_name': f"@{channel}",
                'messages': existing_messages,
                'total_messages': messages_loaded,
                'last_update': datetime.now().isoformat(),
                'status': 'partial',
                'pages_loaded': pages_loaded
            }
            
            with open(temp_cache_file, 'w', encoding='utf-8') as f:
                json.dump(temp_data, f, ensure_ascii=False)
            
            # Определяем, есть ли еще страницы для загрузки
            if has_more and (max_pages is None or pages_loaded < max_pages):
                try:
                    # Устанавливаем переменные для цикла загрузки
                    consecutive_empty_pages = 0
                    retries_left = 5  # Максимальное количество повторных попыток при ошибках
                    max_total_pages = 1000  # Максимальное общее количество страниц для безопасности
                    
                    # Начинаем загрузку остальных страниц
                    while has_more and (max_pages is None or pages_loaded < max_pages) and pages_loaded < max_total_pages:
                        # Определяем ID последнего сообщения для пагинации
                        # Находим сообщение с минимальным ID среди тех, что загрузили на предыдущей странице
                        last_id = None
                        if initial_messages:
                            # Сортируем по ID, чтобы найти минимальный
                            sorted_ids = sorted([msg.get('id') for msg in initial_messages if msg.get('id')])
                            
                            # Проверяем, есть ли ID у первого сообщения (для первой итерации)
                            if pages_loaded == 0 and sorted_ids:
                                last_id = sorted_ids[0]  # Берем минимальный ID для первой итерации
                            else:
                                # Для последующих итераций уменьшаем ID на 20
                                if 'last_used_id' in locals() and last_used_id:
                                    last_id = last_used_id - 20
                                elif sorted_ids:
                                    last_id = sorted_ids[0] - 20  # Берем минимальный ID и уменьшаем на 20
                        
                        # Запоминаем ID для следующей итерации
                        last_used_id = last_id
                        
                        if not last_id:
                            print("Невозможно определить ID для пагинации, завершаем загрузку")
                            break
                        
                        # Делаем паузу перед следующим запросом (случайную, чтобы избежать блокировки)
                        pause_time = 3 + random.random() * 2  # 3-5 секунд
                        print(f"Пауза перед загрузкой страницы {pages_loaded + 1}: {pause_time:.1f} сек")
                        time.sleep(pause_time)
                        
                        # Запрашиваем следующую страницу сообщений
                        print(f"Загрузка страницы {pages_loaded + 1} (используя ID {last_id})...")
                        try:
                            # Вызываем parse_telegram_channel и корректно обрабатываем возвращаемые значения
                            result = parse_telegram_channel(channel, start_id=last_id)
                            if len(result) == 3:
                                new_messages, next_has_more, captcha_html = result
                                if captcha_html:
                                    print(f"Обнаружена CAPTCHA при загрузке страницы {pages_loaded + 1}")
                                    has_more = False
                                    break
                            else:
                                new_messages, next_has_more = result
                            
                            has_more = next_has_more
                            pages_loaded += 1
                            
                            if new_messages and len(new_messages) > 0:
                                # Подсчитываем новые уникальные сообщения
                                new_unique_count = 0
                                for msg in new_messages:
                                    if msg.get('id') and not message_ids.get(msg.get('id')):
                                        existing_messages.append(msg)
                                        message_ids[msg.get('id')] = True
                                        new_unique_count += 1
                                
                                # Обновляем начальное значение для следующей итерации
                                initial_messages = new_messages
                                
                                print(f"Загружена страница {pages_loaded}: получено {len(new_messages)} сообщений, из них {new_unique_count} уникальных")
                                
                                # Сохраняем промежуточный результат с каждым успешным обновлением
                                messages_loaded = len(existing_messages)
                                temp_data = {
                                    'channel': channel,
                                    'channel_name': f"@{channel}",
                                    'messages': existing_messages,
                                    'total_messages': messages_loaded,
                                    'last_update': datetime.now().isoformat(),
                                    'status': 'partial',
                                    'pages_loaded': pages_loaded,
                                    'progress': {
                                        'pages': pages_loaded,
                                        'messages': messages_loaded,
                                        'has_more': has_more
                                    }
                                }
                                
                                with open(temp_cache_file, 'w', encoding='utf-8') as f:
                                    json.dump(temp_data, f, ensure_ascii=False)
                                
                                if new_unique_count > 0:
                                    consecutive_empty_pages = 0  # Сбрасываем счетчик пустых страниц
                                    retries_left = 5  # Восстанавливаем счетчик повторных попыток
                                else:
                                    consecutive_empty_pages += 1
                                    print(f"На странице {pages_loaded} нет новых уникальных сообщений (попытка {consecutive_empty_pages}/3)")
                            else:
                                consecutive_empty_pages += 1
                                print(f"Страница {pages_loaded} не содержит сообщений (попытка {consecutive_empty_pages}/3)")
                            
                            # Если много пустых страниц подряд, вероятно, мы достигли конца
                            if consecutive_empty_pages >= 3:
                                print(f"Получено {consecutive_empty_pages} пустых страниц подряд, завершаем загрузку")
                                has_more = False
                            
                        except Exception as page_error:
                            retries_left -= 1
                            print(f"Ошибка при загрузке страницы {pages_loaded}: {str(page_error)}. Осталось попыток: {retries_left}")
                            
                            # Если закончились попытки, завершаем загрузку
                            if retries_left <= 0:
                                print("Слишком много ошибок подряд, завершаем загрузку")
                                has_more = False
                            
                            # Делаем дополнительную паузу после ошибки
                            time.sleep(5)
                    
                    # Финальная обработка и сортировка сообщений
                    print(f"Всего загружено {len(existing_messages)} сообщений ({pages_loaded} страниц)")
                    
                    # Проверяем на дубликаты еще раз (для надежности)
                    unique_messages = []
                    seen_ids = set()
                    
                    for msg in existing_messages:
                        if msg.get('id') and msg.get('id') not in seen_ids:
                            unique_messages.append(msg)
                            seen_ids.add(msg.get('id'))
                    
                    if len(unique_messages) < len(existing_messages):
                        print(f"Удалено {len(existing_messages) - len(unique_messages)} дубликатов")
                        existing_messages = unique_messages
                    
                    # Сортируем сообщения (новые сначала)
                    existing_messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
                    
                    # Статистика
                    messages_with_images = sum(1 for msg in existing_messages if msg.get('has_images', False))
                    total_messages = len(existing_messages)
                    
                    # Сохраняем итоговый результат
                    final_data = {
                        'channel': channel,
                        'channel_name': f"@{channel}",
                        'messages': existing_messages,
                        'total_messages': total_messages,
                        'messages_with_images': messages_with_images,
                        'last_update': datetime.now().isoformat(),
                        'pages_loaded': pages_loaded,
                        'status': 'complete'
                    }
                    
                    with open(cache_file, 'w', encoding='utf-8') as f:
                        json.dump(final_data, f, ensure_ascii=False)
                    
                    # Удаляем временный файл
                    if os.path.exists(temp_cache_file):
                        os.remove(temp_cache_file)
                    
                    return jsonify({
                        "success": True,
                        "message": f"Загружено и сохранено {total_messages} сообщений (просмотрено {pages_loaded} страниц)",
                        "channel": channel,
                        "count": total_messages,
                        "pages": pages_loaded
                    })
                
                except Exception as e:
                    # В случае ошибки при загрузке
                    import traceback
                    error_details = traceback.format_exc()
                    print(f"Ошибка при загрузке: {str(e)}")
                    print(f"Детали ошибки: {error_details}")
                    
                    # Сохраняем то, что успели загрузить
                    messages_loaded = len(existing_messages)
                    if messages_loaded > 0:
                        try:
                            # Сортируем сообщения по дате (новые сначала)
                            existing_messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
                            
                            # Сохраняем промежуточный результат как финальный
                            partial_data = {
                                'channel': channel,
                                'channel_name': f"@{channel}",
                                'messages': existing_messages,
                                'total_messages': messages_loaded,
                                'last_update': datetime.now().isoformat(),
                                'pages_loaded': pages_loaded,
                                'status': 'partial'  # Отмечаем как частичный результат
                            }
                            
                            with open(cache_file, 'w', encoding='utf-8') as f:
                                json.dump(partial_data, f, ensure_ascii=False)
                            
                            # Удаляем временный файл
                            if os.path.exists(temp_cache_file):
                                os.remove(temp_cache_file)
                            
                            partial_success = True
                        except Exception as save_error:
                            print(f"Ошибка при сохранении промежуточных результатов: {str(save_error)}")
                    
                    if partial_success:
                        return jsonify({
                            "partial_success": True,
                            "message": f"Загрузка была прервана ошибкой, но удалось сохранить {messages_loaded} сообщений из {pages_loaded} страниц",
                            "error": str(e),
                            "channel": channel,
                            "count": messages_loaded,
                            "pages": pages_loaded
                        })
                    else:
                        return jsonify({
                            "success": False,
                            "error": f"Ошибка при загрузке сообщений: {str(e)}",
                            "details": error_details,
                            "channel": channel
                        }), 500
            else:
                # Если нет дополнительных страниц или достигнут лимит страниц
                # Сортируем и сохраняем что уже есть
                existing_messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
                
                final_data = {
                    'channel': channel,
                    'channel_name': f"@{channel}",
                    'messages': existing_messages,
                    'total_messages': len(existing_messages),
                    'last_update': datetime.now().isoformat(),
                    'pages_loaded': pages_loaded,
                    'status': 'complete'
                }
                
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump(final_data, f, ensure_ascii=False)
                
                # Удаляем временный файл, если он есть
                if os.path.exists(temp_cache_file):
                    os.remove(temp_cache_file)
                
                return jsonify({
                    "success": True,
                    "message": f"Загружено и сохранено {len(existing_messages)} сообщений",
                    "channel": channel,
                    "count": len(existing_messages),
                    "pages": pages_loaded
                })
        else:
            # Если первая страница не содержит сообщений
            return jsonify({
                "success": False,
                "error": "Не удалось загрузить сообщения или канал пуст",
                "channel": channel
            })
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Критическая ошибка при загрузке канала: {str(e)}")
        print(f"Детали ошибки: {error_details}")
        
        # Проверяем, есть ли промежуточные результаты
        if os.path.exists(temp_cache_file):
            try:
                with open(temp_cache_file, 'r', encoding='utf-8') as f:
                    temp_data = json.load(f)
                
                if temp_data.get('messages') and len(temp_data['messages']) > 0:
                    # Обновляем статус
                    temp_data['status'] = 'partial'
                    
                    # Сохраняем как основной файл
                    with open(cache_file, 'w', encoding='utf-8') as f:
                        json.dump(temp_data, f, ensure_ascii=False)
                    
                    return jsonify({
                        "partial_success": True,
                        "message": f"Произошла ошибка, но удалось сохранить {len(temp_data['messages'])} сообщений",
                        "error": str(e),
                        "channel": channel,
                        "count": len(temp_data['messages'])
                    })
            except Exception:
                pass
        
        return jsonify({
            "success": False,
            "error": f"Критическая ошибка при загрузке: {str(e)}",
            "details": error_details
        }), 500

def fetch_more_telegram_messages(channel, existing_messages, start_date=None, end_date=None, limit=100, max_pages=None):
    """
    Загружает все доступные сообщения для канала, используя пагинацию.
    
    :param channel: Имя канала без @
    :param existing_messages: Список уже полученных сообщений
    :param start_date: Начальная дата для фильтрации
    :param end_date: Конечная дата для фильтрации
    :param limit: Максимальное количество сообщений для получения на странице
    :param max_pages: Максимальное количество страниц для загрузки
    :return: (список всех сообщений, флаг наличия следующей страницы)
    """
    all_messages = existing_messages.copy()
    page_count = 1
    has_more = True
    
    # Получаем ID последнего сообщения для пагинации
    if existing_messages and len(existing_messages) > 0:
        last_message = existing_messages[-1]
        last_id = last_message.get('id')
        
        # Если нет ID, невозможно загрузить дополнительные сообщения
        if not last_id:
            print("Невозможно определить ID последнего сообщения для пагинации")
            return all_messages, False
        
        # Пока есть еще страницы и не достигнут лимит страниц
        while has_more and (max_pages is None or page_count < max_pages):
            # Пауза между запросами для избежания ограничений API
            print(f"Пауза перед загрузкой следующей страницы {page_count + 1}...")
            time.sleep(3)
            
            print(f"Загрузка страницы {page_count + 1} для канала {channel}, начиная с ID: {last_id}")
            
            # Получаем следующую страницу сообщений
            new_messages, new_has_more = parse_telegram_channel(
                channel, 
                start_date=start_date, 
                end_date=end_date, 
                limit=limit,
                start_id=last_id
            )
            
            # Если не удалось получить сообщения, прерываем цикл
            if not new_messages:
                print(f"Не удалось получить сообщения для страницы {page_count + 1}")
                break
                
            print(f"Получено {len(new_messages)} новых сообщений на странице {page_count + 1}")
            
            # Обновляем флаг наличия следующей страницы
            has_more = new_has_more
            
            # Обновляем ID последнего сообщения для следующей страницы
            if len(new_messages) > 0:
                last_id = new_messages[-1].get('id')
                
                # Если ID не найден, прерываем цикл
                if not last_id:
                    print("Невозможно определить ID последнего сообщения для пагинации")
                    break
                    
                # Добавляем новые сообщения к общему списку
                all_messages.extend(new_messages)
                
            else:
                # Если нет новых сообщений, прерываем цикл
                break
                
            # Увеличиваем счетчик страниц
            page_count += 1
            
            # Если достигнут лимит страниц, останавливаемся
            if max_pages is not None and page_count >= max_pages:
                print(f"Достигнут лимит страниц ({max_pages})")
                break
    
    return all_messages, has_more

@app.route('/api/telegram/refresh_cache', methods=['POST'])
def refresh_telegram_cache():
    """
    Обновляет кэш для указанного канала Telegram
    """
    data = request.json
    channel = data.get('channel', '')
    force_refresh = data.get('force_refresh', True)  # По умолчанию принудительно обновляем
    
    if not channel:
        return jsonify({"status": "error", "error": "Не указан канал Telegram"}), 400
    
    # Удаляем @ из начала, если есть
    if channel.startswith('@'):
        channel = channel[1:]
    
    app.logger.info(f"Запрос обновления кэша для канала {channel} (force_refresh={force_refresh})")
    
    try:
        # Проверяем наличие кэшированных данных
        cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
        cache_file = os.path.join(cache_dir, f"{channel}.json")
        
        existing_messages = []
        existing_ids = set()
        cache_exists = False
        
        # Загружаем существующие сообщения из кэша, если он есть
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                
                if 'messages' in cache_data and isinstance(cache_data['messages'], list):
                    existing_messages = cache_data['messages']
                    existing_ids = {msg.get('id') for msg in existing_messages if msg.get('id')}
                    cache_exists = True
                    
                    app.logger.info(f"Загружено {len(existing_messages)} существующих сообщений из кэша")
            except Exception as e:
                app.logger.error(f"Ошибка при чтении кэша: {str(e)}")
        
        # Запрашиваем новые сообщения у Telegram
        start_time = time.time()
        
        # Если force_refresh=True, игнорируем кэш и делаем новый запрос
        # Иначе возвращаем существующие сообщения из кэша
        if force_refresh or not cache_exists:
            app.logger.info(f"Запрашиваем новые сообщения для канала {channel}")
            
            # Получаем новые сообщения от Telegram
            new_messages, has_more = parse_telegram_channel(channel, force_refresh=True)
            
            # Если получены новые сообщения
            if new_messages:
                # Создаем директорию для кэша, если её нет
                os.makedirs(cache_dir, exist_ok=True)
                
                # Если кэш существует, добавляем только новые сообщения
                if cache_exists:
                    # Смотрим, сколько новых сообщений добавится
                    new_count = 0
                    merged_messages = existing_messages.copy()
                    
                    for msg in new_messages:
                        if msg.get('id') and msg.get('id') not in existing_ids:
                            merged_messages.append(msg)
                            existing_ids.add(msg.get('id'))
                            new_count += 1
                    
                    app.logger.info(f"Добавлено {new_count} новых сообщений к существующим {len(existing_messages)}")
                    
                    # Обновляем кэш только если есть новые сообщения
                    if new_count > 0:
                        try:
                            # Сортируем сообщения по дате (новые сначала)
                            merged_messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
                            
                            # Обновляем файл кэша
                            cache_data = {
                                'channel': channel,
                                'channel_name': f"@{channel}",
                                'messages': merged_messages,
                                'total_messages': len(merged_messages),
                                'last_update': datetime.now().isoformat()
                            }
                            
                            with open(cache_file, 'w', encoding='utf-8') as f:
                                json.dump(cache_data, f, ensure_ascii=False)
                                
                            app.logger.info(f"Кэш для канала {channel} успешно обновлен")
                            
                            # Формируем ответ
                            return jsonify({
                                "status": "success",
                                "message": f"Кэш успешно обновлен, получено {new_count} новых сообщений",
                                "new_count": new_count,
                                "total_count": len(merged_messages),
                                "cache_time": datetime.now().isoformat()
                            })
                        except Exception as e:
                            app.logger.error(f"Ошибка при сохранении обновленного кэша: {str(e)}")
                            return jsonify({
                                "status": "error",
                                "error": f"Ошибка при сохранении кэша: {str(e)}",
                                "cached_messages": len(existing_messages)
                            })
                    else:
                        # Если новых сообщений нет, просто возвращаем информацию
                        return jsonify({
                            "status": "info",
                            "message": "Новых сообщений не найдено",
                            "new_count": 0,
                            "total_count": len(existing_messages),
                            "cache_time": datetime.now().isoformat()
                        })
                # Если кэша нет, создаем новый
                else:
                    try:
                        # Сортируем сообщения по дате (новые сначала)
                        new_messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
                        
                        # Создаем новый файл кэша
                        cache_data = {
                            'channel': channel,
                            'channel_name': f"@{channel}",
                            'messages': new_messages,
                            'total_messages': len(new_messages),
                            'last_update': datetime.now().isoformat()
                        }
                        
                        with open(cache_file, 'w', encoding='utf-8') as f:
                            json.dump(cache_data, f, ensure_ascii=False)
                            
                        app.logger.info(f"Создан новый кэш для канала {channel} с {len(new_messages)} сообщениями")
                        
                        # Формируем ответ
                        return jsonify({
                            "status": "success",
                            "message": f"Создан новый кэш с {len(new_messages)} сообщениями",
                            "new_count": len(new_messages),
                            "total_count": len(new_messages),
                            "cache_time": datetime.now().isoformat()
                        })
                    except Exception as e:
                        app.logger.error(f"Ошибка при создании нового кэша: {str(e)}")
                        return jsonify({
                            "status": "error",
                            "error": f"Ошибка при создании кэша: {str(e)}"
                        })
            else:
                # Если не получены новые сообщения, но есть кэш
                if cache_exists:
                    return jsonify({
                        "status": "info",
                        "message": "Не удалось получить новые сообщения, используется существующий кэш",
                        "cached_messages": len(existing_messages)
                    })
                else:
                    return jsonify({
                        "status": "error",
                        "error": "Не удалось получить сообщения из канала"
                    })
        else:
            # Если не требуется обновление и есть кэш, возвращаем информацию о кэше
            return jsonify({
                "status": "info",
                "message": "Используется существующий кэш",
                "cached_messages": len(existing_messages)
            })
    except Exception as e:
        app.logger.error(f"Ошибка при обновлении кэша: {str(e)}")
        return jsonify({
            "status": "error",
            "error": f"Ошибка при обновлении кэша: {str(e)}"
        })

@app.route('/tzhelper')
@login_required
def tz_helper():
    # Проверяем доступ к функции
    if not has_feature_access(current_user, 'tzhelper'):
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('index'))
    
    # Получаем историю чатов пользователя
    user_chats = TzHelperChat.query.filter_by(user_id=current_user.id).order_by(TzHelperChat.updated_at.desc()).all()
    
    return render_template('tzhelper.html', chats=user_chats)

# API-маршруты для ТЗhelper
@app.route('/api/tzhelper/chats', methods=['GET'])
@login_required
def get_tz_chats():
    """Получение списка чатов пользователя"""
    chats = TzHelperChat.query.filter_by(user_id=current_user.id).order_by(TzHelperChat.updated_at.desc()).all()
    chat_list = []
    
    for chat in chats:
        chat_list.append({
            'id': chat.id,
            'title': chat.title,
            'created_at': chat.created_at.strftime('%d.%m.%Y %H:%M'),
            'updated_at': chat.updated_at.strftime('%d.%m.%Y %H:%M'),
            'has_content': bool(chat.content)
        })
    
    return jsonify({'chats': chat_list})

@app.route('/api/tzhelper/chat/<int:chat_id>', methods=['GET'])
@login_required
def get_tz_chat(chat_id):
    """Получение данных конкретного чата"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    messages = []
    for msg in chat.chat_history:
        messages.append({
            'id': msg.id,
            'role': msg.role,
            'content': msg.content,
            'timestamp': msg.timestamp.strftime('%d.%m.%Y %H:%M:%S')
        })
    
    return jsonify({
        'id': chat.id,
        'title': chat.title,
        'content': chat.content,
        'created_at': chat.created_at.strftime('%d.%m.%Y %H:%M'),
        'updated_at': chat.updated_at.strftime('%d.%m.%Y %H:%M'),
        'messages': messages
    })

@app.route('/api/tzhelper/chat', methods=['POST'])
@login_required
def create_tz_chat():
    """Создание нового чата"""
    title = request.json.get('title', 'Новое ТЗ')
    
    new_chat = TzHelperChat(
        user_id=current_user.id,
        title=title
    )
    
    # Добавляем первое сообщение от ассистента
    initial_message = TzHelperMessage(
        role='assistant',
        content='Здравствуйте! Я помогу вам составить техническое задание. Расскажите, пожалуйста, о вашем проекте или задаче, и я задам уточняющие вопросы для формирования детального ТЗ.'
    )
    
    new_chat.chat_history.append(initial_message)
    
    try:
        db.session.add(new_chat)
        db.session.commit()
        
        # Логируем действие
        log_user_action(f"Создан новый чат ТЗhelper: {title}")
        
        return jsonify({
            'id': new_chat.id,
            'title': new_chat.title,
            'created_at': new_chat.created_at.strftime('%d.%m.%Y %H:%M'),
            'messages': [{
                'id': initial_message.id,
                'role': initial_message.role,
                'content': initial_message.content,
                'timestamp': initial_message.timestamp.strftime('%d.%m.%Y %H:%M:%S')
            }]
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Ошибка при создании чата: {str(e)}'}), 500

@app.route('/api/tzhelper/chat/<int:chat_id>', methods=['PUT'])
@login_required
def update_tz_chat(chat_id):
    """Обновление информации о чате"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    data = request.json
    if 'title' in data:
        chat.title = data['title']
    
    if 'content' in data:
        chat.content = data['content']
    
    try:
        db.session.commit()
        return jsonify({
            'id': chat.id,
            'title': chat.title,
            'content': chat.content,
            'updated_at': chat.updated_at.strftime('%d.%m.%Y %H:%M')
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Ошибка при обновлении чата: {str(e)}'}), 500

@app.route('/api/tzhelper/chat/<int:chat_id>', methods=['DELETE'])
@login_required
def delete_tz_chat(chat_id):
    """Удаление чата"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    try:
        db.session.delete(chat)
        db.session.commit()
        
        # Логируем действие
        log_user_action(f"Удален чат ТЗhelper: {chat.title}")
        
        return jsonify({'success': True}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Ошибка при удалении чата: {str(e)}'}), 500

@app.route('/api/tzhelper/chat/<int:chat_id>/message', methods=['POST'])
@login_required
def send_tz_message(chat_id):
    """Отправка сообщения в чат и получение ответа от API"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    data = request.json
    message_content = data.get('content')
    
    if not message_content or not message_content.strip():
        return jsonify({'error': 'Сообщение не может быть пустым'}), 400
    
    # Создаем сообщение пользователя
    user_message = TzHelperMessage(
        chat_id=chat.id,
        role='user',
        content=message_content
    )
    
    try:
        # Добавляем сообщение пользователя в БД сразу
        db.session.add(user_message)
        db.session.commit()
        
        # Получаем все сообщения в чате для контекста
        chat_messages = TzHelperMessage.query.filter_by(chat_id=chat.id).order_by(TzHelperMessage.created_at).all()
        messages_for_api = [{'role': msg.role, 'content': msg.content} for msg in chat_messages]
        
        # Вызываем API для получения ответа
        print(f"Вызов API для чата {chat_id} с {len(messages_for_api)} сообщениями")
        api_response = call_ai_api(messages_for_api)
        
        # Создаем сообщение от ассистента
        assistant_message = TzHelperMessage(
            chat_id=chat.id,
            role='assistant',
            content=api_response
        )
        
        # Добавляем сообщение от ассистента в БД
        db.session.add(assistant_message)
        
        # Обновляем время последнего взаимодействия с чатом
        chat.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        # Обновляем последнее сообщение в списке чатов
        chat.last_message = api_response[:100] + "..." if len(api_response) > 100 else api_response
        db.session.commit()
        
        # Возвращаем оба сообщения
        return jsonify({
            'user_message': {
                'id': user_message.id,
                'role': user_message.role,
                'content': user_message.content,
                'timestamp': user_message.created_at.isoformat()
            },
            'assistant_message': {
                'id': assistant_message.id,
                'role': assistant_message.role,
                'content': assistant_message.content,
                'timestamp': assistant_message.created_at.isoformat()
            }
        })
        
    except Exception as e:
        # Откатываем транзакцию в случае ошибки
        db.session.rollback()
        print(f"Ошибка при обработке сообщения: {str(e)}")
        
        # Если сообщение пользователя было создано, но произошла ошибка позже,
        # попробуем его сохранить отдельно
        if user_message.id is None:
            try:
                db.session.add(user_message)
                db.session.commit()
            except:
                pass
        
        # Создаем сообщение об ошибке от ассистента
        error_message = f"Извините, произошла ошибка при обработке вашего запроса: {str(e)}"
        
        return jsonify({
            'error': error_message
        }), 500

@app.route('/api/tzhelper/chat/<int:chat_id>/generate-tz', methods=['POST'])
@login_required
def generate_tz(chat_id):
    """Генерация финального ТЗ на основе истории чата"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    # Получаем историю переписки для контекста
    messages_history = []
    for msg in chat.chat_history:
        messages_history.append({
            'role': msg.role,
            'content': msg.content
        })
    
    # Добавляем запрос на генерацию ТЗ
    messages_history.append({
        'role': 'user',
        'content': 'Пожалуйста, сформируй для меня полное техническое задание на основе нашего обсуждения. Структурируй его по разделам, включая цели, требования, функциональность, сроки и другие важные аспекты.'
    })
    
    try:
        # Вызываем API и получаем ответ с ТЗ
        tz_content = call_ai_api(messages_history)
        
        # Создаем сообщение от ассистента с ТЗ
        tz_message = TzHelperMessage(
            chat_id=chat.id,
            role='assistant',
            content=tz_content
        )
        
        # Сохраняем ТЗ в чате
        chat.content = tz_content
        
        # Сохраняем в БД
        db.session.add(tz_message)
        db.session.commit()
        
        # Логируем действие
        log_user_action(f"Сгенерировано ТЗ для чата: {chat.title}")
        
        return jsonify({
            'message': {
                'id': tz_message.id,
                'role': tz_message.role,
                'content': tz_message.content,
                'timestamp': tz_message.timestamp.strftime('%d.%m.%Y %H:%M:%S')
            },
            'tz_content': tz_content
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Ошибка при генерации ТЗ: {str(e)}'}), 500

@app.route('/api/tzhelper/chat/<int:chat_id>/export', methods=['GET'])
@login_required
def export_tz(chat_id):
    """Экспорт ТЗ в документ"""
    chat = TzHelperChat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    if not chat.content:
        return jsonify({'error': 'ТЗ еще не сгенерировано для этого чата'}), 400
    
    try:
        # Создаем временный файл для документа
        from io import BytesIO
        from docx import Document
        
        # Создаем документ
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading(f'Техническое задание: {chat.title}', 0)
        
        # Добавляем информацию о создателе
        doc.add_paragraph(f'Создано: {chat.updated_at.strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Автор: {current_user.first_name} {current_user.last_name}, {current_user.department}')
        
        # Добавляем разделитель
        doc.add_paragraph('-----------------------------------')
        
        # Добавляем содержимое ТЗ
        # Разбиваем текст ТЗ на строки и добавляем их в документ
        for line in chat.content.split('\n'):
            # Проверяем, является ли строка заголовком
            if line.strip().startswith('#'):
                # Определяем уровень заголовка по количеству символов #
                level = line.count('#', 0, line.find(' '))
                heading_text = line.strip('#').strip()
                doc.add_heading(heading_text, level)
            else:
                doc.add_paragraph(line)
        
        # Сохраняем документ в буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Логируем действие
        log_user_action(f"Экспортировано ТЗ для чата: {chat.title}")
        
        # Возвращаем документ для скачивания
        filename = f"TZ_{chat.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Ошибка при экспорте ТЗ: {str(e)}'}), 500

def call_ai_api(messages):
    """Вызов API искусственного интеллекта для получения ответа"""
    try:
        # Конфигурируем работу с API
        use_api = True  # Включаем использование API
        
        # Получаем токен из переменной окружения или используем предоставленный токен
        token = os.environ.get("QWEN_API_TOKEN", "Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpZCI6IjFjODQwNDA3LTk1YzgtNGNlYS1iNjk1LTM1OWQ5YmIxZGYzZSIsImV4cCI6MTc0OTEwNjI3MH0.YQOOlEGN-7dd2YQQrl7lUXz8tQWnJ3LVMtZ0b35Ewx0")
        
        # Удаляем префикс "Bearer " если он присутствует в токене
        if token.startswith("Bearer "):
            token = token[7:]
        
        print(f"Используется токен API: {token[:10]}...")
        
        if use_api:
            # Настройки для API Qwen
            QWEN_URL = "https://chat.qwenlm.ai/api/chat/completions"
            QWEN_HEADERS = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/237.84.2.178 Safari/537.36"
            }
            
            # Преобразуем сообщения в формат, требуемый API Qwen
            qwen_messages = []
            for msg in messages:
                qwen_messages.append({
                    "role": msg["role"],
                    "content": msg["content"],
                    "extra": {},
                    "chat_type": "t2t"
                })
            
            payload = {
                "chat_type": "t2t",
                "messages": qwen_messages,
                "model": "qwen-max-latest",
                "stream": False
            }
            
            # Отправляем запрос к API
            print(f"Отправка запроса к API Qwen с {len(qwen_messages)} сообщениями")
            try:
                response = requests.post(QWEN_URL, headers=QWEN_HEADERS, json=payload, timeout=25)
                
                # Проверяем статус ответа
                if response.status_code == 200:
                    # Получаем ответ от API
                    result = response.json()
                    print(f"Успешный ответ от API Qwen")
                    return result["choices"][0]["message"]["content"]
                else:
                    print(f"Ошибка API: {response.status_code}")
                    print(f"Текст ошибки: {response.text}")
                    # Если ошибка API, используем заглушку
                    error_msg = f"Ошибка API Qwen: {response.status_code}. Текст: {response.text[:100]}... Используется режим демонстрации."
                    print(error_msg)
                    return generate_stub_response(messages)
            except requests.exceptions.RequestException as e:
                print(f"Ошибка соединения с API: {str(e)}")
                # Если ошибка соединения, используем заглушку
                return f"Не удалось соединиться с сервисом ИИ: {str(e)[:100]}... Временно используется режим демонстрации."
        
        # Если use_api=False, используем заглушку
        print("Используем режим демонстрации (заглушку) вместо реального API")
        return generate_stub_response(messages)
            
    except Exception as e:
        # Логируем ошибку
        print(f"Ошибка при вызове AI API: {str(e)}")
        # Возвращаем сообщение об ошибке
        return f"Возникла ошибка при обработке запроса. Временно используется режим демонстрации. Детали ошибки: {str(e)[:100]}..."

def generate_stub_response(messages):
    """Генерирует ответ-заглушку в зависимости от содержания последнего сообщения"""
    # Определяем последнее сообщение пользователя
    last_message = messages[-1]['content'] if messages and messages[-1]['role'] == 'user' else ""
    
    # Формируем контекстный ответ в зависимости от содержания последнего сообщения
    if "новое тз" in last_message.lower() or "создать тз" in last_message.lower() or len(messages) <= 2:
        return """Я помогу вам составить техническое задание. Для начала, расскажите, пожалуйста:
1. Какой продукт или сервис вы хотите разработать?
2. Какие основные функции должны быть в нем? 
3. Для какой целевой аудитории он предназначен?
4. Какие у вас есть предпочтения по технологиям и срокам?"""
    
    elif "генерир" in last_message.lower() or "формир" in last_message.lower() or "готов" in last_message.lower():
        return """На основе нашего обсуждения я подготовил структуру ТЗ. Вот основные разделы, которые мы можем детализировать:

# Техническое задание

## 1. Общее описание проекта
- Название проекта
- Цели и задачи
- Целевая аудитория

## 2. Функциональные требования
- Основной функционал
- Дополнительные возможности
- Интеграции

## 3. Технические требования
- Платформа
- Технологический стек
- Требования к производительности

## 4. Интерфейс и пользовательский опыт
- Общая концепция дизайна
- Структура экранов/страниц
- Макеты и прототипы

## 5. Сроки и этапы разработки
- Основные вехи проекта
- Сроки по задачам

## 6. Бюджет
- Оценка затрат
- График платежей

Хотите ли вы, чтобы я более подробно раскрыл какой-то из разделов?"""
    
    elif "дизайн" in last_message.lower() or "интерфейс" in last_message.lower():
        return """По дизайну и интерфейсу я рекомендую детально прописать в ТЗ следующие аспекты:

1. Общая стилистика: минималистичный, корпоративный, игровой и т.д.
2. Цветовая схема: основные цвета, акценты, градиенты
3. Типографика: шрифты для заголовков и основного текста
4. Компоненты интерфейса: кнопки, поля ввода, таблицы, карточки
5. Адаптивность: как интерфейс должен меняться на разных устройствах
6. Анимации и переходы между экранами
7. Прототипы основных экранов (можно указать, что должны быть предоставлены отдельно)
8. Требования к иконкам и изображениям

Также стоит указать, будет ли дизайн передаваться в Figma, Adobe XD или другом формате, и какие гайдлайны следует соблюдать (например, Material Design, iOS Human Interface Guidelines)."""
    
    else:
        return """Спасибо за предоставленную информацию. Чтобы составить более полное ТЗ, мне нужны дополнительные детали:

1. Какие конкретные функции будут доступны пользователям разных ролей?
2. Требуется ли интеграция с внешними сервисами или API?
3. Какие есть особые требования к безопасности данных?
4. Планируется ли мобильная версия или приложение?
5. Какие метрики успеха проекта вы планируете отслеживать?

Эта информация поможет сделать ТЗ более конкретным и избежать недопонимания на этапе разработки."""

def parse_telegram_messages_from_soup(soup, start_date=None, end_date=None):
    """
    Извлекает сообщения из HTML-страницы канала Telegram
    
    :param soup: объект BeautifulSoup с HTML-страницей
    :param start_date: начальная дата для фильтрации
    :param end_date: конечная дата для фильтрации
    :return: список сообщений
    """
    messages = []
    message_containers = soup.select('.tgme_widget_message_wrap')
    
    app.logger.info(f"Обнаружено {len(message_containers)} сообщений на странице")
    
    # Если передана строковая дата, преобразуем её в объект datetime
    if isinstance(start_date, str):
        try:
            # Предполагаем формат даты ДД.ММ.ГГГГ
            parts = start_date.split('.')
            if len(parts) == 3:
                day, month, year = map(int, parts)
                start_date = datetime(year, month, day)
                app.logger.info(f"Строковая дата start_date преобразована в объект datetime: {start_date}")
        except Exception as e:
            app.logger.error(f"Ошибка при парсинге строковой даты start_date={start_date}: {str(e)}")
            start_date = None
    
    # Аналогично для end_date
    if isinstance(end_date, str):
        try:
            parts = end_date.split('.')
            if len(parts) == 3:
                day, month, year = map(int, parts)
                end_date = datetime(year, month, day)
                app.logger.info(f"Строковая дата end_date преобразована в объект datetime: {end_date}")
        except Exception as e:
            app.logger.error(f"Ошибка при парсинге строковой даты end_date={end_date}: {str(e)}")
            end_date = None
    
    for container in message_containers:
        try:
            # Получаем текстовое содержимое сообщения
            message_text_container = container.select_one('.tgme_widget_message_text')
            
            # Получаем текст сообщения с учетом форматирования HTML
            if message_text_container:
                try:
                    # Сохраняем HTML-версию текста для форматированного отображения
                    # Преобразуем весь контейнер в строку HTML, сохраняя все атрибуты
                    message_html = str(message_text_container)
                    
                    # Проверяем, что html действительно содержит теги
                    has_html_tags = '<' in message_html and '>' in message_html
                    
                    # Получаем текстовую версию для поиска и т.д.
                    message_text = message_text_container.get_text(strip=True)
                    
                    app.logger.debug(f"Найден текст сообщения. Длина текста: {len(message_text)}, длина HTML: {len(message_html)}, содержит HTML-теги: {has_html_tags}")
                    
                    # Если текст пустой или если HTML не содержит тегов, формируем его из текста
                    if (not message_text.strip() or not has_html_tags) and message_html:
                        try:
                            # Если HTML есть, но текст пустой, извлекаем текст еще раз
                            from bs4 import BeautifulSoup
                            html_soup = BeautifulSoup(message_html, 'html.parser')
                            message_text = html_soup.get_text(strip=True)
                            app.logger.debug(f"Текст извлечен из HTML: {len(message_text)} символов")
                            
                            # Если текст всё ещё пустой, но есть HTML
                            if not message_text and has_html_tags:
                                message_text = "Текст сообщения в формате HTML (см. ссылку)"
                        except Exception as e:
                            app.logger.error(f"Ошибка при извлечении текста из HTML: {str(e)}")
                            message_text = "Текст сообщения в формате HTML (см. ссылку)"
                            
                    # Обеспечиваем, что html_content содержит правильный HTML-контент с тегами
                    if not has_html_tags and message_text:
                        # Если HTML не содержит тегов, но есть текст, формируем стандартный HTML-контейнер
                        message_html = f'<div class="tgme_widget_message_text">{message_text}</div>'
                        app.logger.debug(f"Сформирован HTML-контейнер из текста: {len(message_html)} символов")
                except Exception as e:
                    app.logger.error(f"Ошибка при обработке текстового контейнера: {str(e)}")
                    # В случае ошибки сохраняем обработанный текст
                    if message_text and not message_html:
                        message_html = f'<div class="tgme_widget_message_text">{message_text}</div>'
                    else:
                        message_html = '<div class="tgme_widget_message_text">Ошибка при обработке текста</div>'
                        message_text = "Ошибка при обработке текста сообщения"
            else:
                # Если контейнера нет, создаем пустой контейнер с сообщением
                message_html = '<div class="tgme_widget_message_text">Текст сообщения отсутствует</div>'
                message_text = "Текст сообщения отсутствует"
                app.logger.debug("Текстовый контейнер не найден в сообщении")
            
            # Проверяем наличие медиа-контента
            message_photo = container.select_one('.tgme_widget_message_photo_wrap')
            message_video = container.select_one('.tgme_widget_message_video_wrap')
            message_document = container.select_one('.tgme_widget_message_document_wrap')
            
            # Определяем наличие медиа
            has_media = bool(message_photo or message_video or message_document)
            
            # Ограничиваем длину превью текстом в 100 символов
            preview = message_text[:100] + "..." if len(message_text) > 100 else message_text
            
            # Получаем дату сообщения
            date_element = container.select_one('.tgme_widget_message_date time')
            date_timestamp = 0
            if date_element and 'datetime' in date_element.attrs:
                try:
                    datetime_str = date_element['datetime']
                    message_datetime = datetime.fromisoformat(datetime_str.replace('Z', '+00:00'))
                    date_timestamp = message_datetime.timestamp()
                    
                    # Проверяем, попадает ли сообщение в заданный диапазон дат
                    if start_date:
                        # Сравниваем только даты, без времени
                        message_date = message_datetime.date()
                        start_date_date = start_date.date() if isinstance(start_date, datetime) else start_date
                        
                        if message_date < start_date_date:
                            app.logger.debug(f"Сообщение от {message_date} старее {start_date_date}, пропускаем")
                            continue
                    
                    if end_date:
                        # Сравниваем только даты, без времени
                        message_date = message_datetime.date()
                        end_date_date = end_date.date() if isinstance(end_date, datetime) else end_date
                        
                        if message_date > end_date_date:
                            app.logger.debug(f"Сообщение от {message_date} новее {end_date_date}, пропускаем")
                            continue
                    
                    # Форматируем дату для отображения
                    message_date = message_datetime.astimezone(timezone.utc).astimezone()
                except Exception as e:
                    app.logger.error(f"Ошибка при обработке даты сообщения: {str(e)}")
                    # Если не удалось получить дату, используем текущую
                    message_datetime = datetime.now()
                    message_date = message_datetime
            else:
                # Если элемент времени не найден, используем текущую дату
                message_datetime = datetime.now()
                message_date = message_datetime
            
            # Получаем ссылку на сообщение
            link_element = container.select_one('.tgme_widget_message_date')
            message_link = link_element['href'] if link_element and 'href' in link_element.attrs else ""
            
            # Получаем ID сообщения из ссылки
            message_id = None
            match = re.search(r'/(\d+)$', message_link)
            if match:
                message_id = int(match.group(1))
            
            # Форматируем результат
            messages.append({
                'date': message_date.strftime('%d.%m.%Y'),
                'time': message_datetime.strftime('%H:%M:%S'),
                'date_timestamp': date_timestamp,
                'text': message_text,
                'html_content': message_html,  # Добавляем HTML-контент
                'preview': preview,
                'link': message_link,
                'id': message_id,
                'has_images': has_media,
                'images_count': 1 if has_media else 0
            })
        except Exception as e:
            app.logger.error(f"Ошибка при обработке сообщения Telegram: {str(e)}")
            continue
    
    # Сортируем сообщения по убыванию даты (более новые вначале)
    messages.sort(key=lambda x: x.get('date_timestamp', 0), reverse=True)
    
    app.logger.info(f"Извлечено {len(messages)} сообщений после фильтрации")
    
    return messages

def parse_telegram_channel(channel, start_date=None, end_date=None, limit=100, start_id=None, custom_headers=None, force_refresh=False):
    """
    Парсит сообщения из канала Telegram.
    
    :param channel: название канала без @
    :param start_date: начальная дата фильтрации
    :param end_date: конечная дата фильтрации
    :param limit: максимальное количество сообщений
    :param start_id: ID сообщения для пагинации
    :param custom_headers: дополнительные HTTP-заголовки для запроса
    :param force_refresh: принудительное обновление, игнорируя кэш
    :return: (список сообщений, флаг наличия следующей страницы)
    """
    has_more = False
    max_retries = 3
    retry_count = 0
    
    app.logger.info(f"Запрос канала {channel}, start_date={start_date}, end_date={end_date}, limit={limit}, start_id={start_id}, force_refresh={force_refresh}")
    
    # Список разных User-Agent для обхода ограничений
    user_agents = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/94.0.4606.81 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:93.0) Gecko/20100101 Firefox/93.0',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.110 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/97.0.4692.71 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Safari/605.1.15',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.83 Safari/537.36',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.75 Safari/537.36'
    ]
    
    # Проверяем наличие кэшированных данных
    cache_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'telegram_cache')
    cache_file = os.path.join(cache_dir, f"{channel}.json")
    
    # Всегда проверяем кэш
    has_cache = False
    cache_data = None
    if os.path.exists(cache_file):
        try:
            app.logger.info(f"Проверка существующего кэша для канала {channel}")
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            if 'messages' in cache_data and isinstance(cache_data['messages'], list) and cache_data['messages']:
                has_cache = True
                app.logger.info(f"Найден кэш с {len(cache_data['messages'])} сообщениями")
        except Exception as e:
            app.logger.error(f"Ошибка при чтении кэша: {str(e)}")
    
    # Если это запрос для пагинации или принудительное обновление, всегда делаем запрос
    # Иначе просто возвращаем кэш, если он есть
    if not start_id and not force_refresh and has_cache:
        app.logger.info(f"Возвращаем {len(cache_data['messages'])} сообщений из кэша")
        return cache_data['messages'], False
    
    # Функция для обычного запроса
    def make_request(url, headers):
        return requests.get(url, headers=headers, timeout=30)
    
    # Попытка сделать запрос к Telegram
    try:
        # Формируем URL с учетом пагинации (если указан start_id)
        if start_id:
            url = f"https://t.me/s/{channel}?after={start_id}"
        else:
            url = f"https://t.me/s/{channel}"
        
        app.logger.info(f"Запрос к URL: {url}")
        
        # Выбираем случайный User-Agent
        headers = {
            'User-Agent': random.choice(user_agents),
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
            'Cache-Control': 'no-cache',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Referer': 'https://t.me/',
            'Connection': 'keep-alive',
            'Pragma': 'no-cache',
            'DNT': '1'
        }
        
        # Добавляем дополнительные заголовки, если они есть
        if custom_headers:
            headers.update(custom_headers)
            
        # Выполняем HTTP-запрос с заголовками
        response = make_request(url, headers)
        
        # Проверяем код ответа, если не 200, отлавливаем ошибку
        response.raise_for_status()
        
        # Парсим HTML
        app.logger.debug(f"Ответ получен, статус: {response.status_code}")
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Проверяем наличие контейнера с сообщениями
        messages_container = soup.select_one('.tgme_channel_history')
        if not messages_container:
            app.logger.warning(f"Контейнер с сообщениями не найден, использую кэш если доступен")
            
            # Если есть кэш, возвращаем его
            if has_cache:
                return cache_data['messages'], False
            
            # Иначе возвращаем пустой список
            return [], False
            
        # Парсим сообщения из HTML
        messages = parse_telegram_messages_from_soup(soup, start_date, end_date)
        
        # Если есть контейнер для дополнительных сообщений, значит есть еще страницы
        has_more_container = soup.select_one('.tme_messages_more')
        has_more = has_more_container is not None
        
        app.logger.info(f"Успешно получено {len(messages)} сообщений из канала {channel}")
        
        # Сохраняем полученные сообщения в кэш, если это не пагинация
        if not start_id and messages:
            try:
                # Если существует кэш, обновляем его
                if has_cache:
                    # Добавляем новые сообщения к существующим, исключая дубликаты
                    existing_ids = {msg.get('id'): True for msg in cache_data['messages'] if msg.get('id')}
                    new_messages_count = 0
                    
                    for msg in messages:
                        if msg.get('id') and not existing_ids.get(msg.get('id')):
                            cache_data['messages'].append(msg)
                            existing_ids[msg.get('id')] = True
                            new_messages_count += 1
                    
                    # Обновляем дату последнего обновления
                    cache_data['last_update'] = datetime.now().isoformat()
                    
                    app.logger.info(f"Обновлен кэш для канала {channel}, добавлено {new_messages_count} новых сообщений")
                else:
                    # Создаем новый кэш
                    cache_data = {
                        'channel': channel,
                        'channel_name': f"@{channel}",
                        'messages': messages,
                        'total_messages': len(messages),
                        'last_update': datetime.now().isoformat(),
                        'status': 'complete'
                    }
                    
                    app.logger.info(f"Создан новый кэш для канала {channel} с {len(messages)} сообщениями")
                
                # Сохраняем обновленный или новый кэш
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump(cache_data, f, ensure_ascii=False)
            except Exception as e:
                app.logger.error(f"Ошибка при сохранении кэша: {str(e)}")
        
        return messages, has_more
        
    except Exception as e:
        app.logger.error(f"Ошибка при запросе к Telegram: {str(e)}")
        
        # Если это пагинация и ошибка, просто возвращаем пустой список
        if start_id:
            return [], False
            
        # Если есть кэш, возвращаем его в случае ошибки
        if has_cache:
            app.logger.info(f"Использую кэш как запасной вариант, возвращаю {len(cache_data['messages'])} сообщений")
            return cache_data['messages'], False
            
        # Если нет кэша или он пустой
        return [], False

# Импортируем список компаний для поиска
from work_with_bd.companies import COMPANIES

# API ключ для Checko.ru
COMPANY_API_KEY = API_KEY  # Использовать тот же ключ

# Функция для получения данных о компании через API Checko.ru
def get_company_info(query, search_params=None):
    # Базовый URL API
    base_url = "https://api.checko.ru/v2"
    
    try:
        # Поиск компании по запросу с учетом дополнительных параметров
        search_url = f"{base_url}/search?key={API_KEY}"
        
        # Добавляем базовые параметры поиска, если не указаны
        if search_params is None:
            search_params = {
                "by": "name",    # Поиск по наименованию
                "obj": "org",    # Поиск по организациям
                "query": query,  # Текст запроса
                "active": "true" # Только активные организации
            }
        else:
            search_params["query"] = query
            
        # Добавляем все параметры в URL
        for key, value in search_params.items():
            search_url += f"&{key}={value}"
        
        # Выполняем запрос
        search_response = requests.get(search_url)
        search_response.raise_for_status()
        search_data = search_response.json()
        
        # Проверяем, что получили результаты поиска
        if not search_data or "data" not in search_data or not search_data["data"].get("Записи"):
            app.logger.error(f"Компании по запросу '{query}' не найдены")
            return None
        
        # Получаем список компаний из результатов поиска
        companies = search_data["data"]["Записи"]
        
        if not companies:
            app.logger.error(f"Компании по запросу '{query}' не найдены в результатах")
            return None
            
        # Если поиск по названию и есть несколько результатов, ищем точное совпадение
        selected_company = None
        
        if search_params.get("by") == "name" and len(companies) > 0:
            # Ищем точное совпадение по сокращенному наименованию
            for company in companies:
                if "НаимСокр" in company and company["НаимСокр"].lower() == query.lower():
                    selected_company = company
                    break
            
            # Если точного совпадения нет, берем первый результат
            if not selected_company:
                selected_company = companies[0]
        else:
            # Для других типов поиска берем первый результат
            selected_company = companies[0]
        
        # Получаем ИНН компании
        inn = selected_company.get("ИНН")
        
        if not inn:
            app.logger.error("Не удалось получить ИНН компании")
            return None
        
        # Запрос подробной информации о компании
        company_url = f"{base_url}/company?key={API_KEY}&inn={inn}"
        company_response = requests.get(company_url)
        company_response.raise_for_status()
        company_data = company_response.json()
        
        # Запрос финансовой информации
        finance_url = f"{base_url}/company/finance?key={API_KEY}&inn={inn}"
        finance_response = requests.get(finance_url)
        finance_response.raise_for_status()
        finance_data = finance_response.json()
        
        # Сохраняем сырые данные для отображения
        raw_data = {
            "search": search_data,
            "company": company_data,
            "finance": finance_data
        }
        
        # Собираем информацию о компании
        company_info = {
            "name": selected_company.get("НаимСокр", "Нет данных"),
            "full_name": selected_company.get("НаимПолн", "Нет данных"),
            "inn": inn,
            "ogrn": selected_company.get("ОГРН", "Нет данных"),
            "registration_date": selected_company.get("ДатаРег", "Нет данных"),
            "status": selected_company.get("Статус", "Нет данных"),
            "address": selected_company.get("ЮрАдрес", "Нет данных"),
            "ceo": selected_company.get("Руковод", {}).get("ФИО", "Нет данных") if selected_company.get("Руковод") else "Нет данных",
            "activity": selected_company.get("ОКВЭД", "Нет данных"),
            "okved_code": company_data.get("okved", {}).get("code", "Нет данных"),
            "employees_count": company_data.get("employees_count", "Нет данных"),
            "financial_data": {},
            "raw_data": raw_data,  # Сохраняем сырые данные
            "finance_raw": finance_data  # Для отображения детальной финансовой отчетности
        }
        
        # Обрабатываем финансовые данные
        if finance_data and "data" in finance_data and "years" in finance_data["data"]:
            app.logger.info(f"Обрабатываем финансовые данные из API 1")
            for year_data in finance_data["data"]["years"]:
                year = year_data.get("year")
                revenue = year_data.get("Выручка")
                profit = year_data.get("ЧистПрибыль")
                
                if year:
                    if revenue is not None:
                        company_info["financial_data"][f"Выручка {year}"] = f"{revenue:,} ₽".replace(",", " ")
                    if profit is not None:
                        company_info["financial_data"][f"Прибыль {year}"] = f"{profit:,} ₽".replace(",", " ")
        
        # Если не нашли финансовые данные в первом API, пробуем использовать альтернативный
        if not company_info["financial_data"] and finance_data and "years" in finance_data:
            app.logger.info(f"Не найдены данные в API 1, используем API 2")
            for year_data in finance_data["years"]:
                year = year_data.get("year")
                revenue = year_data.get("revenue")
                profit = year_data.get("profit")
                
                if year:
                    if revenue is not None:
                        company_info["financial_data"][f"Выручка {year}"] = f"{revenue:,} ₽".replace(",", " ")
                    if profit is not None:
                        company_info["financial_data"][f"Прибыль {year}"] = f"{profit:,} ₽".replace(",", " ")
            
            # Сохраняем альтернативные финансовые данные, если они есть
            company_info["finance_raw"] = finance_data
        
        app.logger.info(f"Финансовые данные после обработки: {company_info['financial_data']}")
        
        return company_info
    
    except Exception as e:
        app.logger.error(f"Ошибка при получении данных: {str(e)}")
        traceback.print_exc()
        return None

# Функция для получения данных о компании по ИНН
def get_company_by_inn(inn):
    base_url = "https://api.checko.ru/v2"
    headers = {
        "Authorization": f"Token {API_KEY}"
    }
    
    # URL для получения данных компании по ИНН
    company_url = f"{base_url}/company/inn/{inn}"
    finances_url = f"{base_url}/finances/inn/{inn}/years"
    
    company_data = {}
    finance_data = {}
    
    try:
        # Запрашиваем данные о компании
        company_response = requests.get(company_url, headers=headers)
        company_response.raise_for_status()
        company_data = company_response.json()
        
        # Сохраняем для отладки
        if debug_mode:
            print(f"Данные компании {inn}:", company_data)
    except Exception as company_error:
        if debug_mode:
            print(f"Ошибка при получении данных компании: {str(company_error)}")
        return {"error": f"Ошибка при получении данных компании: {str(company_error)}"}
    
    try:
        # Запрашиваем финансовые данные
        finance_response = requests.get(finances_url, headers=headers)
        finance_response.raise_for_status()
        finance_data = finance_response.json()
        
        if debug_mode:
            print(f"Финансовые данные {inn}:", finance_data)
    except Exception as finance_error:
        if debug_mode:
            print(f"Ошибка при получении финансовых данных: {str(finance_error)}")
        # Не возвращаем ошибку, т.к. финансовые данные не являются критичными
    
    # Сохраняем сырые данные для отображения и отладки
    raw_data = {
        "company": company_data,
        "finance": finance_data
    }
    
    # Получаем основные данные о компании
    data = company_data.get("data", {})
    
    # Подробное логирование для отладки структуры данных
    if debug_mode:
        print("Структура данных компании:", list(data.keys()))
        if "ОКВЭД" in data:
            print("Структура основного ОКВЭД:", data["ОКВЭД"])
        if "ОКВЭДДоп" in data:
            print(f"Количество дополнительных ОКВЭД: {len(data.get('ОКВЭДДоп', []))}")
    
    # Получаем информацию об адресе
    address = data.get("ЮрАдрес", "Нет данных")
    if isinstance(address, dict):
        address_data = address
    else:
        address_data = "Нет данных"
    
    # Получаем информацию об уставном капитале
    authorized_capital = None
    if "УстКап" in data:
        authorized_capital = {
            "type": data["УстКап"].get("Тип", "Нет данных"),
            "amount": data["УстКап"].get("Сумма", 0)
        }
    
    # Получаем информацию об основном ОКВЭД
    main_okved = None
    if "ОКВЭД" in data:
        main_okved = {
            "code": data["ОКВЭД"].get("Код", "Нет данных"),
            "name": data["ОКВЭД"].get("Наим", "Нет данных"),
            "version": data["ОКВЭД"].get("Версия", "Нет данных")
        }
        if debug_mode:
            print(f"Основной ОКВЭД: {main_okved}")

    # Получаем дополнительные ОКВЭД с расширенной проверкой
    additional_okved = []
    if "ОКВЭДДоп" in data:
        if isinstance(data["ОКВЭДДоп"], list):
            for okved in data["ОКВЭДДоп"]:
                if isinstance(okved, dict):
                    additional_okved.append({
                        "code": okved.get("Код", "Нет данных"),
                        "name": okved.get("Наим", "Нет данных"),
                        "version": okved.get("Версия", "Нет данных")
                    })
        if debug_mode:
            print(f"Найдено {len(additional_okved)} дополнительных ОКВЭД")
    
    # Собираем информацию о компании
    company_info = {
        "name": data.get("НаимСокр", "Нет данных"),
        "full_name": data.get("НаимПолн", "Нет данных"),
        "inn": inn,
        "ogrn": data.get("ОГРН", "Нет данных"),
        "registration_date": data.get("ДатаРег", "Нет данных"),
        "status": data.get("Статус", "Нет данных"),
        "address": address_data,
        "ceo": "Нет данных",
        "activity": main_okved["name"] if main_okved else data.get("ОКВЭД", "Нет данных"),
        "okved_code": main_okved["code"] if main_okved else "Нет данных",
        "okved_version": main_okved["version"] if main_okved else "Нет данных",
        "additional_okved": additional_okved,
        "authorized_capital": authorized_capital,
        "employees_count": data.get("КолРаб", "Нет данных"),
        "financial_data": {},
        "raw_data": raw_data,
        "finance_raw": finance_data
    }

    if "Руковод" in data and isinstance(data["Руковод"], list) and len(data["Руковод"]) > 0:
        company_info["ceo"] = data["Руковод"][0].get("ФИО", "Нет данных")

    # Обработка финансовых данных (основные показатели)
    if finance_data and "data" in finance_data and "years" in finance_data["data"]:
        for year_data in finance_data["data"]["years"]:
            year = year_data.get("year")
            revenue = year_data.get("Выручка")
            profit = year_data.get("ЧистПрибыль")
            if year:
                if revenue is not None:
                    company_info["financial_data"][f"Выручка {year}"] = f"{revenue:,} ₽".replace(",", " ")
                if profit is not None:
                    company_info["financial_data"][f"Прибыль {year}"] = f"{profit:,} ₽".replace(",", " ")

    return company_info

# Функция для получения описания кода строки финансовой отчетности
def get_financial_code_description(code):
    # Словарь с описаниями основных кодов строк
    descriptions = {
        "1100": "Внеоборотные активы",
        "1110": "Нематериальные активы",
        "1120": "Результаты исследований и разработок",
        "1150": "Основные средства",
        "1200": "Оборотные активы",
        "1210": "Запасы",
        "1230": "Дебиторская задолженность",
        "1240": "Финансовые вложения",
        "1250": "Денежные средства",
        "1300": "Капитал и резервы",
        "1310": "Уставный капитал",
        "1350": "Добавочный капитал",
        "1360": "Резервный капитал",
        "1370": "Нераспределенная прибыль (непокрытый убыток)",
        "1400": "Долгосрочные обязательства",
        "1410": "Заемные средства",
        "1500": "Краткосрочные обязательства",
        "1510": "Заемные средства",
        "1520": "Кредиторская задолженность",
        "2110": "Выручка",
        "2120": "Себестоимость продаж",
        "2100": "Валовая прибыль (убыток)",
        "2200": "Прибыль (убыток) от продаж",
        "2300": "Прибыль (убыток) до налогообложения",
        "2400": "Чистая прибыль (убыток)"
    }
    
    return descriptions.get(code, "")

@app.route('/companies')
@login_required
def companies():
    # Проверяем доступ к функции
    if not has_feature_access(current_user, "companies"):
        flash('У вас нет доступа к этой функции', 'danger')
        return redirect(url_for('dashboard'))
    
    # Логируем действие пользователя
    log_user_action("Доступ к модулю поиска информации по компаниям")
    
    return render_template('companies.html', title='Поиск информации по компаниям')

@app.route('/api/company/search', methods=['POST'])
@login_required
def search_company():
    if not has_feature_access(current_user, "companies"):
        return jsonify({"error": "Нет доступа к этой функции"}), 403
    
    data = request.get_json()
    if not data:
        return jsonify({"error": "Неверный формат данных"}), 400
    
    query = data.get('query')
    search_by = data.get('search_by', 'name')  # по умолчанию поиск по наименованию
    
    if not query:
        return jsonify({"error": "Необходимо указать поисковый запрос"}), 400
    
    app.logger.info(f"Поиск компании: {query}, поиск по: {search_by}")
    
    # Настройка параметров поиска
    search_params = {"by": search_by, "obj": "org", "active": "true"}
    
    try:
        # Если поиск по ИНН, используем другую функцию
        if search_by == "inn":
            company_info = get_company_by_inn(query)
        else:
            company_info = get_company_info(query, search_params)
        
        if not company_info:
            return jsonify({"error": f"Компания по запросу '{query}' не найдена"}), 404
        
        # Логируем действие
        log_user_action(f"Поиск информации о компании: {company_info.get('name')} (ИНН: {company_info.get('inn')})")
        
        # Возвращаем результаты поиска
        return jsonify({
            "success": True,
            "company": company_info
        })
    except Exception as e:
        app.logger.error(f"Ошибка при поиске компании: {str(e)}")
        return jsonify({"error": f"Произошла ошибка при поиске: {str(e)}"}), 500

@app.route('/api/company/inn/<inn>', methods=['GET'])
@login_required
def get_company_info_by_inn(inn):
    if not has_feature_access(current_user, "companies"):
        return jsonify({"error": "Нет доступа к этой функции"}), 403
    
    app.logger.info(f"Запрос информации о компании по ИНН: {inn}")
    
    company_info = get_company_by_inn(inn)
    
    if not company_info:
        return jsonify({"error": f"Компания с ИНН '{inn}' не найдена"}), 404
    
    # Логируем действие
    log_user_action(f"Получение информации о компании: {company_info.get('name')} (ИНН: {inn})")
    
    # Возвращаем результаты
    return jsonify({
        "success": True,
        "company": company_info
    })

@app.route('/api/company/list', methods=['GET'])
@login_required
def get_company_list():
    if not has_feature_access(current_user, "companies"):
        return jsonify({"error": "Нет доступа к этой функции"}), 403
    
    try:
        # Импортируем список компаний из файла
        from work_with_bd.companies import COMPANIES
        
        # Логируем действие
        log_user_action("Запрос списка компаний")
        
        # Преобразуем словарь в список словарей для удобства работы на клиенте
        companies_list = [{"name": name, "inn": inn} for name, inn in COMPANIES.items()]
        
        # Сортируем по названию компании
        companies_list.sort(key=lambda x: x["name"])
        
        return jsonify({
            "success": True,
            "companies": companies_list
        })
    except Exception as e:
        app.logger.error(f"Ошибка при получении списка компаний: {str(e)}")
        return jsonify({"error": f"Произошла ошибка при получении списка компаний: {str(e)}"}), 500

if __name__ == '__main__':
    # Создаем необходимую структуру каталогов
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(PROFILE_PICS_DIR, exist_ok=True)
    os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)
    
    # Запускаем скрипт миграции базы данных
    try:
        from work_with_bd.db_migrate import migrate_database
        migrate_database()
    except Exception as e:
        print(f"Ошибка при миграции базы данных: {e}")
    
    # Запускаем проверку сессий при старте сервера
    # Это позволит обновить статистику и проверить IP адреса
    try:
        with app.app_context():
            print("Инициализация приложения...")
            from work_with_bd.fix_access import fix_sessions
            fix_sessions()
    except Exception as e:
        print(f"Ошибка при инициализации приложения: {e}")
    
    # Запускаем сервер в режиме отладки
    app.run(debug=True, host='0.0.0.0', port=5000)